import streamlit as st
import pandas as pd
from openpyxl import Workbook, load_workbook
from openpyxl.styles import PatternFill, Font, Alignment, Border, Side
from openpyxl.utils import get_column_letter
import io
import re

from docx import Document
from docx.shared import Pt, RGBColor, Twips
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

st.set_page_config(page_title="דוח סקר קרקע", layout="wide", page_icon="🧪")
st.title("🧪 מערכת עיבוד תוצאות מעבדה")
st.caption("v3.8 - RTL sheets, number formatting, compound names left-aligned")
st.markdown("---")

YELLOW_FILL   = PatternFill("solid", fgColor="FFFF00")
ORANGE_FILL   = PatternFill("solid", fgColor="FFC000")
HDR_BLUE_FILL = PatternFill("solid", fgColor="B7D7F0")
HDR_CYAN_FILL = PatternFill("solid", fgColor="00B0F0")

def thin_border():
    s = Side(style="thin", color="000000")
    return Border(left=s, right=s, top=s, bottom=s)

def style_hdr(cell, fill=None, sz=11):
    cell.font      = Font(bold=True, name="Arial", size=sz)
    cell.alignment = Alignment(horizontal="center", vertical="center", wrap_text=True)
    cell.border    = thin_border()
    if fill: cell.fill = fill

def fmt_number(val):
    """הוסף פסיקים למספרים מעל 999"""
    if val is None or val == "": return val
    s = str(val).strip()
    if s.startswith("<") or s.startswith(">"): return s
    try:
        f = float(s)
        if abs(f) < 1000: return val
        if "." in s:
            decimals = len(s.split(".")[-1])
            return f"{f:,.{decimals}f}"
        return f"{int(f):,}"
    except (ValueError, TypeError):
        return val

def style_data(cell, hl=None, sz=10, left_align=False):
    cell.font      = Font(bold=bool(hl), name="Arial", size=sz)
    cell.alignment = Alignment(horizontal="left" if left_align else "center",
                               vertical="center", wrap_text=True)
    cell.border    = thin_border()
    if   hl == "tier1": cell.fill = ORANGE_FILL
    elif hl == "vsl":   cell.fill = YELLOW_FILL
    if cell.value is not None:
        cell.value = fmt_number(cell.value)

def norm(s):
    s = "" if s is None else str(s).strip().lower()
    return re.sub(r"\s+", " ", s.replace("\xa0", " "))

def to_float(v):
    s = str(v).strip() if v is not None else ""
    try: return float(s.lstrip("<>").strip())
    except: return None

def sort_key(sid):
    m = re.match(r"S-?(\d+)", str(sid), re.I)
    return int(m.group(1)) if m else 9999

def parse_sample(sname):
    sname = str(sname).strip()
    if "DUP" in sname.upper(): return None, None
    # S85 (0.5) או S1 (1.0)
    m = re.match(r"^(S\d+[A-Za-z0-9]*)\s*\(([0-9.]+)\)", sname)
    if m: return m.group(1), float(m.group(2))
    # S1-1.0
    m = re.match(r"^(S\d+)-([0-9]+\.?[0-9]*)$", sname)
    if m: return m.group(1), float(m.group(2))
    # 24.15 (3.0) — שם קידוח מספרי עם נקודה
    m = re.match(r"^(\d+[\.\d]*)\s*\(([0-9.]+)\)", sname)
    if m: return m.group(1), float(m.group(2))
    # שם קידוח בלי עומק — קבל כמות שהוא
    return sname, None

def check_exceed(val_str, vsl, tier1):
    if not val_str or str(val_str).strip().startswith("<"): return None
    f = to_float(val_str)
    if f is None: return None
    try:
        t1f = float(tier1) if (tier1 is not None and str(tier1) not in ("-","NA","") and pd.notna(tier1)) else None
        vf  = float(vsl)   if (vsl   is not None and str(vsl)   not in ("-","NA","") and pd.notna(vsl))   else None
        if t1f and t1f > 0 and f > t1f: return "tier1"
        if vf  and vf  > 0 and f > vf:  return "vsl"
    except: pass
    return None

def apply_sid_merge(ws, sid_rows, col=1):
    for sid, rows_list in sid_rows.items():
        if len(rows_list) > 1:
            ws.merge_cells(start_row=rows_list[0], start_column=col,
                           end_row=rows_list[-1], end_column=col)
            c = ws.cell(rows_list[0], col)
            c.alignment = Alignment(horizontal="center", vertical="center", wrap_text=True)
            c.border = thin_border()

METAL_MAP = {
    "aluminium":"Al","aluminum":"Al","antimony":"Sb","arsenic":"As",
    "barium":"Ba","beryllium":"Be","bismuth":"Bi","boron":"B",
    "cadmium":"Cd","calcium":"Ca","chromium":"Cr","cobalt":"Co",
    "copper":"Cu","iron":"Fe","lead":"Pb","lithium":"Li",
    "magnesium":"Mg","manganese":"Mn","mercury":"Hg","nickel":"Ni",
    "potassium":"K","selenium":"Se","silver":"Ag","sodium":"Na",
    "vanadium":"V","zinc":"Zn","molybdenum":"Mo","tin":"Sn",
    "titanium":"Ti","strontium":"Sr","thallium":"Tl",
    "phosphorus":"P","sulphur":"S","silicon":"Si",
}

METALS_ORDER = [
    "Al","Sb","As","Ba","Be","Bi","B","Cd","Ca","Cr","Co","Cu","Fe",
    "Pb","Li","Mg","Mn","Hg","Ni","K","Se","Ag","Na","V","Zn"
]

THRESH_METAL_MAP = {
    "aluminum":"Al","antimony (metallic)":"Sb","antimony":"Sb",
    "arsenic, inorganic":"As","arsenic":"As","barium":"Ba",
    "beryllium and compounds":"Be","beryllium":"Be",
    "boron and borates only":"B","boron":"B",
    "cadmium (water) source: water and air":"Cd","cadmium":"Cd",
    "calcium":"Ca","chromium, total":"Cr","chromium":"Cr","cobalt":"Co",
    "copper":"Cu","iron":"Fe","lead and compounds":"Pb","lead":"Pb",
    "lithium":"Li","magnesium":"Mg","manganese (non-diet)":"Mn","manganese":"Mn",
    "mercuric chloride (and other mercury salts)":"Hg","mercury":"Hg",
    "nickel soluble salts":"Ni","nickel":"Ni","potassium":"K",
    "selenium":"Se","silver":"Ag","sodium":"Na",
    "vanadium and compounds":"V","vanadium":"V",
    "zinc and compounds":"Zn","zinc":"Zn","molybdenum":"Mo","tin":"Sn",
    "titanium":"Ti","strontium":"Sr","thallium":"Tl",
    "phosphorus":"P","sulphur":"S","silicon":"Si",
}

VOC_COMPOUND_ORDER = [
    ("VOCs","Non-Halogenated VOCs","1.2.4-Trimethylbenzene"),
    ("VOCs","Non-Halogenated VOCs","1.3.5-Trimethylbenzene"),
    ("VOCs","Non-Halogenated VOCs","MTBE"),
    ("VOCs","Non-Halogenated VOCs","Styrene"),
    ("VOCs","Non-Halogenated VOCs","n-Butylbenzene"),
    ("VOCs","Non-Halogenated VOCs","n-Propylbenzene"),
    ("VOCs","Non-Halogenated VOCs","Isopropylbenzene"),
    ("VOCs","Non-Halogenated VOCs","Acetone"),
    ("VOCs","Non-Halogenated VOCs","2-Butanone (MEK)"),
    ("VOCs","Non-Halogenated VOCs","1.4-Dioxane"),
    ("VOCs","BTEX","Benzene"),
    ("VOCs","BTEX","Toluene"),
    ("VOCs","BTEX","Ethylbenzene"),
    ("VOCs","BTEX","Sum of Xylenes"),
    ("VOCs","Halogenated VOCs","1.1-Dichloroethane"),
    ("VOCs","Halogenated VOCs","1.1-Dichloroethene"),
    ("VOCs","Halogenated VOCs","1.2-Dichloroethane"),
    ("VOCs","Halogenated VOCs","1.2-Dichloropropane"),
    ("VOCs","Halogenated VOCs","Chlorobenzene"),
    ("VOCs","Halogenated VOCs","Chloroform"),
    ("VOCs","Halogenated VOCs","Dichloromethane"),
    ("VOCs","Halogenated VOCs","Tetrachloroethene"),
    ("VOCs","Halogenated VOCs","Tetrachloromethane"),
    ("VOCs","Halogenated VOCs","Trichloroethene"),
    ("VOCs","Halogenated VOCs","Vinyl chloride"),
    ("VOCs","Halogenated VOCs","cis-1.2-Dichloroethene"),
    ("VOCs","Halogenated VOCs","trans-1.2-Dichloroethene"),
    ("VOCs","Halogenated VOCs","1.4-Dichlorobenzene"),
    ("VOCs","Halogenated VOCs","1.2-Dichlorobenzene"),
    ("VOCs","Halogenated VOCs","1.3-Dichlorobenzene"),
    ("SVOCs","Phenols & Naphtols","2.4-Dimethylphenol"),
    ("SVOCs","Phenols & Naphtols","2-Methylphenol"),
    ("SVOCs","Phenols & Naphtols","3 & 4-Methylphenol"),
    ("SVOCs","Phenols & Naphtols","4-Chloro-3-methylphenol"),
    ("SVOCs","Phenols & Naphtols","Phenol"),
    ("SVOCs","PAHs","Acenaphthene"),
    ("SVOCs","PAHs","Acenaphthylene"),
    ("SVOCs","PAHs","Anthracene"),
    ("SVOCs","PAHs","Benz(a)anthracene"),
    ("SVOCs","PAHs","Benzo(a)pyrene"),
    ("SVOCs","PAHs","Benzo(b)fluoranthene"),
    ("SVOCs","PAHs","Benzo(g.h.i)perylene"),
    ("SVOCs","PAHs","Benzo(k)fluoranthene"),
    ("SVOCs","PAHs","Chrysene"),
    ("SVOCs","PAHs","Dibenz(a.h)anthracene"),
    ("SVOCs","PAHs","Fluoranthene"),
    ("SVOCs","PAHs","Fluorene"),
    ("SVOCs","PAHs","Indeno(1.2.3.cd)pyrene"),
    ("SVOCs","PAHs","Naphthalene"),
    ("SVOCs","PAHs","Phenanthrene"),
    ("SVOCs","PAHs","Pyrene"),
    ("SVOCs","Anilines","4-Chloroaniline"),
    ("SVOCs","Anilines","Aniline"),
    ("SVOCs","Anilines","Benzidine"),
    ("SVOCs","Anilines","Diphenylamine"),
    ("SVOCs","Aromatic Compounds","1,1'-Biphenyl"),
    ("SVOCs","Aromatic Compounds","1-Chloronaphthalene"),
    ("SVOCs","Aromatic Compounds","2-Chloronaphthalene"),
    ("SVOCs","Aromatic Compounds","2-Methylnaphthalene"),
    ("SVOCs","Aromatic Compounds","4-Bromophenyl phenyl ether"),
    ("SVOCs","Aromatic Compounds","4-Chlorophenyl phenyl ether"),
    ("SVOCs","Aromatic Compounds","Carbazole"),
    ("SVOCs","Aromatic Compounds","Dibenzofuran"),
    ("SVOCs","Alcohols","Benzyl Alcohol"),
    ("SVOCs","Aldehydes / Ketones","6-Caprolactam"),
    ("SVOCs","Aldehydes / Ketones","Acetophenone"),
    ("SVOCs","Aldehydes / Ketones","Isophorone"),
    ("SVOCs","Chlorophenols","2-Chlorophenol"),
    ("SVOCs","Chlorophenols","2.4.5-Trichlorophenol"),
    ("SVOCs","Chlorophenols","2.4.6-Trichlorophenol"),
    ("SVOCs","Chlorophenols","2.4-Dichlorophenol"),
    ("SVOCs","Chlorophenols","2.6-Dichlorophenol"),
    ("SVOCs","Chlorophenols","Pentachlorophenol"),
    ("SVOCs","Nitroaromatic Compounds","2.4-Dinitrophenol"),
    ("SVOCs","Nitroaromatic Compounds","2.4-Dinitrotoluene"),
    ("SVOCs","Nitroaromatic Compounds","2-Nitroaniline"),
    ("SVOCs","Nitroaromatic Compounds","2-Nitrophenol"),
    ("SVOCs","Nitroaromatic Compounds","2.6-Dinitrotoluene"),
    ("SVOCs","Nitroaromatic Compounds","3-Nitroaniline"),
    ("SVOCs","Nitroaromatic Compounds","4.6-Dinitro-2-methylphenol"),
    ("SVOCs","Nitroaromatic Compounds","4-Nitroaniline"),
    ("SVOCs","Nitroaromatic Compounds","4-Nitrophenol"),
    ("SVOCs","Nitroaromatic Compounds","Nitrobenzene"),
    ("SVOCs","Chlorinated Hydrocarbons","Bis(2-chloroethoxy)methane"),
    ("SVOCs","Chlorinated Hydrocarbons","Bis(2-chloroethyl)ether"),
    ("SVOCs","Chlorinated Hydrocarbons","Bis(2-chloroisopropyl)ether"),
    ("SVOCs","Nitrosoamines","N-Nitrosodi-n-propylamine"),
    ("SVOCs","Pesticides","Dinoseb"),
    ("SVOCs","Phthalates","Bis(2-ethylhexyl)phthalate"),
    ("SVOCs","Phthalates","Butyl benzyl phthalate"),
    ("SVOCs","Phthalates","Di-n-butyl phthalate"),
    ("SVOCs","Phthalates","Di-n-octyl phthalate"),
    ("SVOCs","Phthalates","Diethyl phthalate"),
    ("SVOCs","Phthalates","Dimethyl phthalate"),
]

# ── VOC ALIAS ─────────────────────────────────────────────────────────────────────
VOC_ALIAS = {
    "1.2.4-trimethylbenzene":           "Trimethylbenzene, 1,2,4-",
    "1,2,4-trimethylbenzene":           "Trimethylbenzene, 1,2,4-",
    "1.3.5-trimethylbenzene":           "Trimethylbenzene, 1,3,5-",
    "1,3,5-trimethylbenzene":           "Trimethylbenzene, 1,3,5-",
    "mtbe":                             "Methyl tert-Butyl Ether (MTBE)",
    "methyl tert-butyl ether":          "Methyl tert-Butyl Ether (MTBE)",
    "n-propylbenzene":                  "Propyl benzene",
    "propylbenzene":                    "Propyl benzene",
    "isopropylbenzene":                 "Cumene",
    "cumene":                           "Cumene",
    "2-butanone (mek)":                 "Methyl Ethyl Ketone - MEK (2-Butanone)",
    "2-butanone":                       "Methyl Ethyl Ketone - MEK (2-Butanone)",
    "methyl ethyl ketone":              "Methyl Ethyl Ketone - MEK (2-Butanone)",
    "mek":                              "Methyl Ethyl Ketone - MEK (2-Butanone)",
    "sum of xylenes":                   "Xylenes",
    "xylenes, total":                   "Xylenes",
    "total xylenes":                    "Xylenes",
    "1.1-dichloroethene":               "1,1-Dichloroethylene",
    "1,1-dichloroethene":               "1,1-Dichloroethylene",
    "1,1-dichloroethylene":             "1,1-Dichloroethylene",
    "1.2-dichloroethane":               "1,2- (EDC) Dichloroethane",
    "1,2-dichloroethane":               "1,2- (EDC) Dichloroethane",
    "ethylene dichloride":              "1,2- (EDC) Dichloroethane",
    "dichloromethane":                  "Methylene Chloride",
    "methylene chloride":               "Methylene Chloride",
    "tetrachloroethene":                "Tetrachloroethylene (PCE)",
    "tetrachloroethylene":              "Tetrachloroethylene (PCE)",
    "perchloroethylene":                "Tetrachloroethylene (PCE)",
    "pce":                              "Tetrachloroethylene (PCE)",
    "tetrachloromethane":               "Carbon Tetrachloride",
    "carbon tetrachloride":             "Carbon Tetrachloride",
    "trichloroethene":                  "Trichloroethylene (TCE)",
    "trichloroethylene":                "Trichloroethylene (TCE)",
    "tce":                              "Trichloroethylene (TCE)",
    "cis-1.2-dichloroethene":           "1,2-cis-Dichloroethylene",
    "cis-1,2-dichloroethene":           "1,2-cis-Dichloroethylene",
    "cis-1,2-dichloroethylene":         "1,2-cis-Dichloroethylene",
    "1,2-cis-dichloroethylene":         "1,2-cis-Dichloroethylene",
    "trans-1.2-dichloroethene":         "1,2-trans-Dichloroethylene",
    "trans-1,2-dichloroethene":         "1,2-trans-Dichloroethylene",
    "trans-1,2-dichloroethylene":       "1,2-trans-Dichloroethylene",
    "1,2-trans-dichloroethylene":       "1,2-trans-Dichloroethylene",
    "benz(a)anthracene":                "Benz[a]anthracene",
    "benz[a]anthracene":                "Benz[a]anthracene",
    "benzo(g.h.i)perylene":             "h,i)perylene Benzo(g",
    "benzo(g,h,i)perylene":             "h,i)perylene Benzo(g",
    "benzo[g,h,i]perylene":             "h,i)perylene Benzo(g",
    "dibenz(a.h)anthracene":            "h]anthracene Dibenz[a",
    "dibenz(a,h)anthracene":            "h]anthracene Dibenz[a",
    "dibenz[a,h]anthracene":            "h]anthracene Dibenz[a",
    "indeno(1.2.3.cd)pyrene":           "2,3-cd]pyrene Indeno[1",
    "indeno(1,2,3-cd)pyrene":           "2,3-cd]pyrene Indeno[1",
    "indeno[1,2,3-cd]pyrene":           "2,3-cd]pyrene Indeno[1",
    "4-chloroaniline":                  "p-Chloroaniline",
    "p-chloroaniline":                  "p-Chloroaniline",
    "1-chloronaphthalene":              "Beta-Chloronaphthalene",
    "2-chloronaphthalene":              "Beta-Chloronaphthalene",
    "6-caprolactam":                    "Caprolactam",
    "2.4.5-trichlorophenol":            "Trichlorophenol, 2,4,5-",
    "2,4,5-trichlorophenol":            "Trichlorophenol, 2,4,5-",
    "2.4.6-trichlorophenol":            "Trichlorophenol, 2,4,6-",
    "2,4,6-trichlorophenol":            "Trichlorophenol, 2,4,6-",
    "2.6-dichlorophenol":               "2,6-Dimethylphenol",
    "2,6-dichlorophenol":               "2,6-Dimethylphenol",
    "4.6-dinitro-2-methylphenol":       "4,6-Dinitro-o-cresol",
    "4,6-dinitro-2-methylphenol":       "4,6-Dinitro-o-cresol",
    "3-nitroaniline":                   "3,5-Dinitroaniline",
    "bis(2-chloroisopropyl)ether":      "Bis(2-chloro-1-methylethyl) ether",
    "n-nitrosodi-n-propylamine":        "N-Nitroso-di-N-propylamine",
    "di-n-butyl phthalate":             "Dibutyl Phthalate",
    "butyl benzyl phthalate":           "Butyl Benzyl Phthalate",
    "di-n-octyl phthalate":             "di-N-Octyl Phthalate",
    "diethyl phthalate":                "Diethyl Phthalate",
    "dimethyl phthalate":               "Dimethylterephthalate",
    "4-chloro-3-methylphenol":          "p-chloro-m-Cresol",
    "1.4-dioxane":                      "1,4-Dioxane",
    "1.2-dichlorobenzene":              "1,2-Dichlorobenzene",
    "1.4-dichlorobenzene":              "1,4-Dichlorobenzene",
    "1.1-dichloroethane":               "1,1-Dichloroethane",
    "1.2-dichloropropane":              "1,2-Dichloropropane",
    "2.4-dimethylphenol":               "2,4-Dimethylphenol",
    "2.4-dichlorophenol":               "2,4-Dichlorophenol",
    "2.4-dinitrophenol":                "2,4-Dinitrophenol",
    "2.4-dinitrotoluene":               "2,4-Dinitrotoluene",
    "2.6-dinitrotoluene":               "2,6-Dinitrotoluene",
    "bis(2-ethylhexyl)phthalate":       "Bis(2-ethylhexyl)phthalate",
    "1,1'-biphenyl":                   "1,1'-Biphenyl",
    "n-butylbenzene":                   "n-Butylbenzene",
    "vinyl chloride":                   "Vinyl Chloride",
}

# ── PFAS ALIAS ────────────────────────────────────────────────────────────────────
PFAS_ALIAS = {
    "2,3,3,3-tetrafluoro-2-(heptafluoropropoxy)propanoic acid (hfpo-da)":
        "hexafluoropropylene oxide dimer acid (hfpo-da)",
    "7h-perfluoroheptanoic acid (hpfhpa)":      "perfluoroheptanoic acid (pfhpa)",
    "perfluorobutane sulfonic acid (pfbs)":     "perfluorobutanesulfonic acid (pfbs)",
    "perfluorobutane sulfonate (pfbs)":         "perfluorobutanesulfonic acid (pfbs)",
    "perfluorohexane sulfonic acid (pfhxs)":    "perfluorohexanesulfonic acid (pfhxs)",
    "perfluorohexane sulfonate (pfhxs)":        "perfluorohexanesulfonic acid (pfhxs)",
    "perfluorooctane sulfonic acid (pfos)":     "perfluorooctanesulfonic acid (pfos)",
    "perfluorooctane sulfonate (pfos)":         "perfluorooctanesulfonic acid (pfos)",
    "perfluorooctadecanoic acid (pfocda)":      "perfluorooctadecanoic acid (pfoda)",
    "perfluoroundecanoic acid (pfunda)":        "perfluoroundecanoic acid (pfuda)",
    "perfluorotetradecanoic acid (pfcpda)":     "perfluorotetradecanoic acid (pfteta)",
    "perfluorodecane sulfonic acid (pfds)":     "perfluorodecanesulfonic acid (pfds)",
    "perfluoroheptane sulfonic acid (pfhps)":   "perfluoroheptanesulfonic acid (pfhps)",
    "perfluoropentane sulfonic acid (pfpes)":   "perfluoropentanesulfonic acid (pfpes)",
    "perfluorooctane sulfonamide (fosa)":       "perfluorooctanesulfonamide (fosa)",
    "perfluoropentanoic acid (pfpea)":          "perfluoropentanoic acid (pfpea)",
    "perfluorodecanoic acid (pfda)":            "perfluorodecanoic acid (pfda)",
    "perfluorododecanoic acid (pfdoda)":        "perfluorododecanoic acid (pfdoda)",
    "perfluoroheptanoic acid (pfhpa)":          "perfluoroheptanoic acid (pfhpa)",
    "perfluorotridecanoic acid (pftrda)":       "perfluorotridecanoic acid (pftrda)",
    "perfluorooctanesulfonic acid (pfos)":      "perfluorooctanesulfonic acid (pfos)",
}

CANONICAL_KEY = "__CANONICAL_MAP_INTERNAL__"

def canonical_compound(name: str) -> str:
    s = norm(name)
    s = s.replace("ethylene", "ethen").replace("ethene", "ethen")
    s = re.sub(r"\s*\([^)]+\)\s*$", "", s)
    nums = re.findall(r"\d+", s)
    num_part = ",".join(nums) if nums else ""
    base_no_nums = re.sub(r"[0-9.,/\-]", " ", s)
    base_no_nums = re.sub(r"\s+", "", base_no_nums)
    canon = (num_part + " " + base_no_nums).strip()
    return canon

def match_threshold(compound_name, thresh_dict):
    key = norm(compound_name)
    canon_map = thresh_dict.get(CANONICAL_KEY)

    for k in (key, key.replace(".",","), key.replace(",",".")):
        if k in thresh_dict and k != CANONICAL_KEY: return thresh_dict[k]

    # VOC alias
    aliased_voc = VOC_ALIAS.get(key) or VOC_ALIAS.get(key.replace(".",","))
    if aliased_voc:
        a_key = norm(aliased_voc)
        for k in (a_key, a_key.replace(".",","), a_key.replace(",",".")):
            if k in thresh_dict and k != CANONICAL_KEY: return thresh_dict[k]

    # PFAS alias
    aliased = PFAS_ALIAS.get(key)
    if aliased:
        a_key = norm(aliased)
        for k in (a_key, a_key.replace(".",","), a_key.replace(",",".")):
            if k in thresh_dict and k != CANONICAL_KEY: return thresh_dict[k]

    stripped = re.sub(r"\s*\([A-Z0-9:_\-]+\)\s*$", "", compound_name).strip().lower()
    for k, v in thresh_dict.items():
        if k == CANONICAL_KEY: continue
        k_s = re.sub(r"\s*\([A-Z0-9:_\-]+\)\s*$", "", k).strip().lower()
        if len(stripped) > 8 and stripped == k_s: return v

    if canon_map:
        ck = canonical_compound(compound_name)
        hit = canon_map.get(ck)
        if hit: return hit

    return {}

def load_threshold_file(file_bytes):
    wb = load_workbook(io.BytesIO(file_bytes), data_only=True)
    ws = wb.active
    thresh = {}
    for row in ws.iter_rows(min_row=6, values_only=True):
        if not row[0]: continue
        name = str(row[0]).strip()
        cas  = str(row[1]).strip() if row[1] else "-"
        def g(ci):
            return (row[ci] if ci < len(row) and row[ci] is not None
                    and str(row[ci]) not in ("NA","") else None)
        thresh[norm(name)] = {
            "name": name, "cas": cas,
            "units": str(row[3]) if row[3] else "mg/kg",
            "VSL": g(4), "Ind_A_06": g(8), "Ind_A_6p": g(9),
            "Ind_B": g(10), "Res_A_06": g(11), "Res_A_6p": g(12), "Res_B": g(13),
        }
    canon_map = {}
    for k, v in thresh.items():
        ck = canonical_compound(k)
        canon_map.setdefault(ck, v)
    thresh[CANONICAL_KEY] = canon_map
    return thresh

def get_tier1_col(land_use, aquifer, depth):
    ind = "industrial" in land_use.lower()
    b   = "b-1" in aquifer.lower()
    if b: return "Ind_B" if ind else "Res_B"
    deep = ">6" in depth
    if ind: return "Ind_A_06" if not deep else "Ind_A_6p"
    else:   return "Res_A_06" if not deep else "Res_A_6p"

def tier1_label(land_use, aquifer, depth):
    return f"TIER 1\n{land_use}\n{aquifer}\n{depth}"

def get_thresh(compound, thresh_dict, t1col):
    t = match_threshold(compound, thresh_dict)
    return t.get("VSL"), t.get(t1col), t.get("cas", "-")

def build_metals_thresh(thresh_dict, t1col):
    result = {}
    for key, v in thresh_dict.items():
        if key == CANONICAL_KEY: continue
        sym = THRESH_METAL_MAP.get(key)
        if sym and sym not in result:
            result[sym] = {"vsl": v.get("VSL"), "tier1": v.get(t1col), "cas": v.get("cas","-")}
    return result

def parse_als_file(file_bytes, filename):
    """
    Parser גנרי - עובד עם כל קבצי ALS:
    - מחפש Client Sample ID לפי תוכן (לא לפי מיקום שורה)
    - מחפש Parameter לפי תוכן
    - מזהה אוטומטית עמודות Unit, LOR, VSL
    - עובד עם PFAS / VOC / Metals / TPH באותו קוד
    """
    try:
        wb = load_workbook(io.BytesIO(file_bytes), data_only=True)
    except Exception as e:
        return None, str(e)

    main = next((wb[n] for n in wb.sheetnames if "Client" in n and "SOIL" in n), wb.worksheets[0])
    rows = list(main.iter_rows(values_only=True))

    # ── מצא שורת Client Sample ID ──────────────────────────────────────────────
    sid_row_idx = next(
        (i for i,r in enumerate(rows)
         if any("Client Sample ID" in str(v) for v in r if v)), None)
    if sid_row_idx is None:
        return None, "לא נמצאה שורת Client Sample ID"

    sid_row = rows[sid_row_idx]
    # מצא את העמודה שמכילה "Client Sample ID"
    sid_label_col = next(ci for ci,v in enumerate(sid_row) if v and "Client Sample ID" in str(v))
    # כל העמודות אחריה = שמות הדגימות
    col2sample = {
        ci: str(v).strip()
        for ci,v in enumerate(sid_row)
        if ci > sid_label_col and v and str(v).strip() not in ("", "None")
    }

    # ── מצא שורת Parameter ─────────────────────────────────────────────────────
    ph_idx = next(
        (i for i,r in enumerate(rows) if r and r[0] == "Parameter"), None)
    if ph_idx is None:
        return None, "לא נמצאה שורת Parameter"

    param_row = rows[ph_idx]

    # זהה אוטומטית עמודות Unit ו-LOR לפי כותרת (לא לפי מיקום קשיח)
    unit_col = next(
        (ci for ci,v in enumerate(param_row)
         if v and str(v).strip().lower() == "unit"), 2)
    lor_col = next(
        (ci for ci,v in enumerate(param_row)
         if v and str(v).strip().lower() == "lor"), unit_col + 1)

    # ── קרא נתונים ─────────────────────────────────────────────────────────────
    records = []
    group = "Unknown"

    for row in rows[ph_idx + 1:]:
        p = row[0] if len(row) > 0 else None
        if not p or str(p).strip() in ("", "None"):
            continue

        # שורת קבוצה: יש שם אבל אין Method (עמודה 1)
        method = row[1] if len(row) > 1 else None
        if not method or str(method).strip() in ("", "None"):
            group = str(p).strip()
            continue

        u   = row[unit_col] if unit_col < len(row) else None
        lor = row[lor_col]  if lor_col  < len(row) else None

        for ci, sname in col2sample.items():
            sid, depth_val = parse_sample(sname)
            if sid is None:
                continue
            val = row[ci] if ci < len(row) else None
            rs  = str(val).strip() if val is not None else ""
            if rs in ("", "None"):
                continue
            result = None
            if rs.startswith("<"):
                result = 0.0
            else:
                try: result = float(rs)
                except: result = None
            if result is not None:
                lor_val = None
                if rs.startswith("<"):
                    try: lor_val = float(rs[1:].strip())
                    except: lor_val = 0.0
                records.append({
                    "sample_id":      sid,
                    "depth":          depth_val,
                    "compound":       str(p).strip(),
                    "compound_lower": norm(p),
                    "unit":           str(u).strip() if u else "mg/kg",
                    "lor":            lor,
                    "result":         result,
                    "result_str":     rs,
                    "lor_val":        lor_val,
                    "group":          group,
                    "source":         filename,
                })

    if not records:
        return None, "לא נמצאו נתונים"
    return pd.DataFrame(records), None

# ── TPH SHEET ─────────────────────────────────────────────────────────────────────
def write_tph_sheet(ws, df, thresh_dict, t1col, t1lbl):
    def is_dro(c):
        if "(dro)" in c or "- dro" in c or c.strip()=="dro": return True
        if "(oro)" in c or "- oro" in c: return False
        if "c10" in c and "c28" in c and "c40" not in c: return True
        return False
    def is_oro(c):
        if "(oro)" in c or "- oro" in c or c.strip()=="oro": return True
        if "(dro)" in c or "- dro" in c: return False
        if "c24" in c and "c40" in c: return True
        if "c28" in c and "c40" in c: return True
        return False
    def is_total(c):
        if any(x in c for x in ["(dro)","(oro)","- dro","- oro"]): return False
        if "c10" in c and "c40" in c: return True
        if "total" in c and ("tph" in c or "hydrocarbon" in c): return True
        return False

    vsl_d,t1_d,_ = get_thresh("C10 - C28 Fraction (DRO)", thresh_dict, t1col)
    vsl_o,t1_o,_ = get_thresh("C24 - C40 Fraction (ORO)", thresh_dict, t1col)
    vsl_t,t1_t,_ = get_thresh("TPH - DRO + ORO (Tier 1)", thresh_dict, t1col)
    vv = [v for v in [vsl_d,vsl_o,vsl_t] if v]
    tt = [v for v in [t1_d,t1_o,t1_t] if v]
    vsl_tot = min(vv) if vv else 350
    t1_tot  = min(tt) if tt else 350

    for ci,h in enumerate(["שם קידוח","עומק","TPH DRO","TPH ORO","Total TPH"],1):
        style_hdr(ws.cell(1,ci,h),HDR_BLUE_FILL)
    ws.merge_cells(start_row=1,start_column=1,end_row=5,end_column=1)
    c=ws.cell(1,1,"שם קידוח"); c.font=Font(bold=True,name="Arial",size=11)
    c.fill=HDR_BLUE_FILL; c.alignment=Alignment(horizontal="center",vertical="center",wrap_text=True); c.border=thin_border()
    sub_rows=["יחידות","CAS","VSL",t1lbl]
    sub_vals={"יחידות":"mg/kg","CAS":"C10-C40","VSL":vsl_tot,t1lbl:t1_tot}
    for ri,lbl in enumerate(sub_rows,2):
        style_hdr(ws.cell(ri,2,lbl),HDR_BLUE_FILL)
        for ci in [3,4,5]: style_hdr(ws.cell(ri,ci,sub_vals[lbl]),HDR_BLUE_FILL)

    pivoted = {}
    for _,r in df.iterrows():
        k=(r["sample_id"],r["depth"])
        if k not in pivoted: pivoted[k]={"DRO":"","ORO":"","TOT":"","DRO_f":None,"ORO_f":None,"DRO_lor":None,"ORO_lor":None}
        c=r["compound_lower"]
        if is_dro(c) and not pivoted[k]["DRO"]:
            pivoted[k]["DRO"]=r["result_str"]; pivoted[k]["DRO_f"]=r["result"]; pivoted[k]["DRO_lor"]=r.get("lor_val")
        elif is_oro(c) and not pivoted[k]["ORO"]:
            pivoted[k]["ORO"]=r["result_str"]; pivoted[k]["ORO_f"]=r["result"]; pivoted[k]["ORO_lor"]=r.get("lor_val")
        elif is_total(c) and not pivoted[k]["TOT"]:
            pivoted[k]["TOT"]=r["result_str"]

    ri=6; prev_sid=None; sid_rows={}
    for (sid,depth_val),v in sorted(pivoted.items(),key=lambda x:(sort_key(x[0][0]),x[0][1] or 0)):
        if v["TOT"]: total_s=v["TOT"]
        else:
            dro_lor=v["DRO"] and str(v["DRO"]).startswith("<")
            oro_lor=v["ORO"] and str(v["ORO"]).startswith("<")
            dro_empty=not v["DRO"]; oro_empty=not v["ORO"]
            dro_num=(v["DRO_lor"] if dro_lor and v["DRO_lor"] is not None else (v["DRO_f"] or 0))
            oro_num=(v["ORO_lor"] if oro_lor and v["ORO_lor"] is not None else (v["ORO_f"] or 0))
            total_f=dro_num+oro_num
            if (dro_lor or dro_empty) and (oro_lor or oro_empty) and not (dro_empty and oro_empty):
                total_s=f"<{total_f:.0f}"
            else: total_s=f"{total_f:.0f}"
        hl_dro=check_exceed(v["DRO"],vsl_tot,t1_tot)
        hl_oro=check_exceed(v["ORO"],vsl_tot,t1_tot)
        hl_total=check_exceed(total_s,vsl_tot,t1_tot)
        if sid!=prev_sid: sid_rows[sid]=[]
        sid_rows[sid].append(ri)
        sid_val=sid if sid!=prev_sid else None; prev_sid=sid
        style_data(ws.cell(ri,1,sid_val)); style_data(ws.cell(ri,2,depth_val))
        style_data(ws.cell(ri,3,v["DRO"]),hl_dro); style_data(ws.cell(ri,4,v["ORO"]),hl_oro)
        style_data(ws.cell(ri,5,total_s),hl_total); ri+=1
    apply_sid_merge(ws,sid_rows,col=1)
    for col,w in zip("ABCDE",[14,10,16,16,16]): ws.column_dimensions[col].width=w
    ws.row_dimensions[1].height=15; ws.freeze_panes="A6"
    ws.sheet_view.rightToLeft=True

# ── METALS SHEET ──────────────────────────────────────────────────────────────────
def write_metals_sheet(ws, df, thresh_dict, t1col, t1lbl):
    df=df.copy(); df["sym"]=df["compound_lower"].map(METAL_MAP); df=df[df["sym"].notna()]
    if df.empty: ws.cell(1,1,"אין נתוני מתכות"); return
    present=set(df["sym"].unique())
    metals=[m for m in METALS_ORDER if m in present]+sorted(present-set(METALS_ORDER))
    mt=build_metals_thresh(thresh_dict,t1col)
    ws.merge_cells(start_row=1,start_column=1,end_row=5,end_column=1)
    c=ws.cell(1,1,"שם קידוח"); c.font=Font(bold=True,name="Arial",size=11)
    c.fill=HDR_BLUE_FILL; c.alignment=Alignment(horizontal="center",vertical="center",wrap_text=True); c.border=thin_border()
    for ci,h in enumerate(["עומק"]+metals,2): style_hdr(ws.cell(1,ci,h),HDR_BLUE_FILL)
    for ri,lbl in enumerate(["יחידות","CAS","VSL",t1lbl],2):
        style_hdr(ws.cell(ri,2,lbl),HDR_BLUE_FILL)
        for ci,sym in enumerate(metals,3):
            t=mt.get(sym,{})
            val={"יחידות":"mg/kg","CAS":t.get("cas","-"),"VSL":t.get("vsl","-"),t1lbl:t.get("tier1","-")}.get(lbl,"-")
            style_hdr(ws.cell(ri,ci,val),HDR_BLUE_FILL)
    pt=df.pivot_table(index=["sample_id","depth"],columns="sym",values="result_str",aggfunc="first")
    pt=pt.reindex(sorted(pt.index,key=lambda x:(sort_key(x[0]),x[1] or 0)))
    ri=6; prev_sid=None; sid_rows={}
    for (sid,depth_val),row_data in pt.iterrows():
        if sid!=prev_sid: sid_rows[sid]=[]
        sid_rows[sid].append(ri)
        sid_val=sid if sid!=prev_sid else None; prev_sid=sid
        style_data(ws.cell(ri,1,sid_val)); style_data(ws.cell(ri,2,depth_val))
        for ci,sym in enumerate(metals,3):
            val=row_data.get(sym,"") or ""
            val="" if str(val)=="nan" else str(val)
            hl=check_exceed(val,mt.get(sym,{}).get("vsl"),mt.get(sym,{}).get("tier1"))
            style_data(ws.cell(ri,ci,val),hl)
        ri+=1
    apply_sid_merge(ws,sid_rows,col=1)
    ws.column_dimensions["A"].width=14; ws.column_dimensions["B"].width=10
    for ci in range(3,len(metals)+3): ws.column_dimensions[get_column_letter(ci)].width=11
    ws.freeze_panes="C6"
    ws.sheet_view.rightToLeft=True

# ── PFAS SHEET ────────────────────────────────────────────────────────────────────
def write_pfas_sheet(ws, df, thresh_dict, t1col, t1lbl):
    df=df.copy()
    pairs_pfas=sorted(df[["sample_id","depth"]].drop_duplicates().values.tolist(),
                      key=lambda x:(sort_key(x[0]),-(x[1] or 0)))
    def to_ug(v):
        if v is None: return None
        try: return round(float(v)*1000,6)
        except: return v
    fixed_hdrs=["שם התרכובת","CAS","VSL [µg/kg]",t1lbl,"יחידות"]
    for ci,h in enumerate(fixed_hdrs,1):
        style_hdr(ws.cell(1,ci,h),HDR_BLUE_FILL)
        ws.merge_cells(start_row=1,start_column=ci,end_row=2,end_column=ci)
        ws.cell(1,ci).alignment=Alignment(horizontal="center",vertical="center",wrap_text=True)
        ws.cell(1,ci).fill=HDR_BLUE_FILL; ws.cell(1,ci).border=thin_border()
    style_hdr(ws.cell(1,6,"LOR"),HDR_BLUE_FILL); style_hdr(ws.cell(2,6,"[µg/kg]"),HDR_BLUE_FILL)
    style_hdr(ws.cell(1,7,"שם קידוח"),HDR_BLUE_FILL); style_hdr(ws.cell(2,7,"עומק"),HDR_BLUE_FILL)
    prev_sid=None; sid_merge_start_p={}
    for ci,(sid,depth_val) in enumerate(pairs_pfas,8):
        sid_val=sid if sid!=prev_sid else None
        style_hdr(ws.cell(1,ci,sid_val),HDR_BLUE_FILL)
        style_hdr(ws.cell(2,ci,depth_val),HDR_BLUE_FILL)
        if sid!=prev_sid: sid_merge_start_p[sid]=ci
        prev_sid=sid
    for sid,start_ci in sid_merge_start_p.items():
        cols_p=[ci for ci,(s,_) in enumerate(pairs_pfas,8) if s==sid]
        if len(cols_p)>1:
            ws.merge_cells(start_row=1,start_column=start_ci,end_row=1,end_column=cols_p[-1])
            c=ws.cell(1,start_ci); c.alignment=Alignment(horizontal="center",vertical="center")
            c.fill=HDR_BLUE_FILL; c.border=thin_border()
    for row_i,cmp in enumerate(df["compound"].unique(),3):
        df_c=df[df["compound"]==cmp]
        vsl_mg,tier1_mg,cas=get_thresh(cmp,thresh_dict,t1col)
        vsl=to_ug(vsl_mg); tier1=to_ug(tier1_mg)
        unit=df_c.iloc[0]["unit"] if not df_c.empty else "µg/kg"
        lor=df_c.iloc[0]["lor"] if not df_c.empty else ""
        style_data(ws.cell(row_i,1,cmp),left_align=True)
        for ci,val in enumerate([cas,vsl,tier1,unit],2): style_data(ws.cell(row_i,ci,val))
        ws.merge_cells(start_row=row_i,start_column=6,end_row=row_i,end_column=6)
        style_data(ws.cell(row_i,6,lor)); style_data(ws.cell(row_i,7,None))
        for ci,(sid,depth_val) in enumerate(pairs_pfas,8):
            sub=df_c[(df_c["sample_id"]==sid)&(df_c["depth"]==depth_val)]
            rs=sub.iloc[0]["result_str"] if not sub.empty else ""
            style_data(ws.cell(row_i,ci,rs),check_exceed(rs,vsl,tier1))
    ws.column_dimensions["A"].width=50
    for ci in range(2,8): ws.column_dimensions[get_column_letter(ci)].width=13
    for ci in range(8,8+len(pairs_pfas)): ws.column_dimensions[get_column_letter(ci)].width=12
    ws.row_dimensions[1].height=60
    ws.freeze_panes="H3"
    ws.sheet_view.rightToLeft=True

# ── VOC+SVOC SHEET ────────────────────────────────────────────────────────────────
def write_voc_sheet(ws, df, thresh_dict, t1col, t1lbl):
    df=df.copy()
    pairs=sorted(df[["sample_id","depth"]].drop_duplicates().values.tolist(),
                 key=lambda x:(sort_key(x[0]),-(x[1] or 0)))
    # Headers A-F: merge rows 1-2
    for ci,h in enumerate(["קבוצה","קבוצה","שם התרכובת","CAS","VSL",t1lbl],1):
        ws.merge_cells(start_row=1,start_column=ci,end_row=2,end_column=ci)
        style_hdr(ws.cell(1,ci,h),HDR_BLUE_FILL,sz=9)
        ws.cell(1,ci).alignment=Alignment(horizontal="center",vertical="center",wrap_text=True)
        ws.cell(1,ci).fill=HDR_BLUE_FILL; ws.cell(1,ci).border=thin_border()
    # G:H merged = יחידות
    ws.merge_cells(start_row=1,start_column=7,end_row=2,end_column=8)
    style_hdr(ws.cell(1,7,"יחידות"),HDR_BLUE_FILL,sz=9)
    ws.cell(1,7).alignment=Alignment(horizontal="center",vertical="center",wrap_text=True)
    ws.cell(1,7).fill=HDR_BLUE_FILL; ws.cell(1,7).border=thin_border()
    # col I: שם קידוח / עומק
    style_hdr(ws.cell(1,9,"שם קידוח"),HDR_BLUE_FILL,sz=9)
    style_hdr(ws.cell(2,9,"עומק"),HDR_BLUE_FILL,sz=9)
    # sample columns from col 10
    prev_sid=None; sid_col_start={}
    for ci,(sid,depth_val) in enumerate(pairs,10):
        sid_val=sid if sid!=prev_sid else None
        style_hdr(ws.cell(1,ci,sid_val),HDR_BLUE_FILL,sz=9)
        style_hdr(ws.cell(2,ci,depth_val),HDR_BLUE_FILL,sz=9)
        if sid!=prev_sid: sid_col_start[sid]=ci
        prev_sid=sid
    for sid,sc in sid_col_start.items():
        cols=[ci for ci,(s,_) in enumerate(pairs,10) if s==sid]
        if len(cols)>1:
            ws.merge_cells(start_row=1,start_column=sc,end_row=1,end_column=cols[-1])
            c=ws.cell(1,sc); c.alignment=Alignment(horizontal="center",vertical="center")
            c.fill=HDR_BLUE_FILL; c.border=thin_border()
    # ALS lookup
    als_data={}
    for _,r in df.iterrows():
        k=norm(r["compound"]); als_data.setdefault(k,{})[(r["sample_id"],r["depth"])]=r["result_str"]
    for k in list(als_data.keys()):
        for alt in (k.replace(".",","),k.replace(",",".")):
            if alt not in als_data: als_data[alt]=als_data[k]
    # data rows
    for row_i,(vs,grp,cmp) in enumerate(VOC_COMPOUND_ORDER,3):
        vsl,tier1,cas=get_thresh(cmp,thresh_dict,t1col)
        cmp_key=norm(cmp)
        cmp_data=als_data.get(cmp_key) or als_data.get(cmp_key.replace(".",",")) or {}
        style_data(ws.cell(row_i,1,vs),sz=9); style_data(ws.cell(row_i,2,grp),sz=9)
        style_data(ws.cell(row_i,3,cmp),sz=9,left_align=True); style_data(ws.cell(row_i,4,cas),sz=9)
        style_data(ws.cell(row_i,5,vsl),sz=9); style_data(ws.cell(row_i,6,tier1),sz=9)
        ws.merge_cells(start_row=row_i,start_column=7,end_row=row_i,end_column=8)
        c=ws.cell(row_i,7,"mg/kg"); c.font=Font(name="Arial",size=9)
        c.alignment=Alignment(horizontal="center",vertical="center"); c.border=thin_border()
        style_data(ws.cell(row_i,9,None),sz=9)
        for ci,(sid,depth_val) in enumerate(pairs,10):
            rs=cmp_data.get((sid,depth_val),"")
            style_data(ws.cell(row_i,ci,rs),check_exceed(rs,vsl,tier1),sz=9)
    # merge col A: VOCs rows 3-32, SVOCs rows 33-96
    for r1,r2,val in [(3,32,"VOCs"),(33,96,"SVOCs")]:
        ws.merge_cells(start_row=r1,start_column=1,end_row=r2,end_column=1)
        c=ws.cell(r1,1,val); c.font=Font(bold=True,name="Arial",size=9)
        c.alignment=Alignment(horizontal="center",vertical="center",wrap_text=True); c.border=thin_border()
    # merge col B
    b_ranges=[
        (3,12,"Non-Halogenated VOCs"),(13,16,"BTEX"),(17,32,"Halogenated VOCs"),
        (33,37,"Phenols & Naphtols"),(38,53,"PAHs"),(54,57,"Anilines"),
        (58,65,"Aromatic Compounds"),(66,66,"Alcohols"),(67,69,"Aldehydes / Ketones"),
        (70,75,"Chlorophenols"),(76,85,"Nitroaromatic Compounds"),
        (86,88,"Chlorinated Hydrocarbons"),(89,89,"Nitrosoamines"),
        (90,90,"Pesticides"),(91,96,"Phthalates"),
    ]
    for r1,r2,val in b_ranges:
        if r2>r1: ws.merge_cells(start_row=r1,start_column=2,end_row=r2,end_column=2)
        c=ws.cell(r1,2,val); c.font=Font(bold=True,name="Arial",size=9)
        c.alignment=Alignment(horizontal="center",vertical="center",wrap_text=True); c.border=thin_border()
    ws.column_dimensions["A"].width=8; ws.column_dimensions["B"].width=22
    ws.column_dimensions["C"].width=35; ws.column_dimensions["D"].width=12
    ws.column_dimensions["E"].width=10; ws.column_dimensions["F"].width=12
    ws.column_dimensions["G"].width=7;  ws.column_dimensions["H"].width=7
    ws.column_dimensions["I"].width=12
    for ci in range(10,10+len(pairs)): ws.column_dimensions[get_column_letter(ci)].width=10
    ws.row_dimensions[1].height=60; ws.row_dimensions[2].height=15
    ws.freeze_panes="J3"
    ws.sheet_view.rightToLeft=True


# ── WORD EXPORT FUNCTIONS ────────────────────────────────────────────────────


# ── צבעים ────────────────────────────────────────────────────────────────────
COLOR_YELLOW = "FFFF00"
COLOR_ORANGE = "FFC000"
COLOR_HEADER = "B7D7F0"
COLOR_HEADER2 = "00B0F0"
COLOR_WHITE   = "FFFFFF"
COLOR_LEGEND_BG = "F2F2F2"

# ── גדלי דף (EMU: 1 inch = 914400) ──────────────────────────────────────────
PAGE_SIZES = {
    # (width_portrait, height_portrait) in Twips (1/20 pt, 1 inch = 1440 twips)
    "A4":      (11906, 16838),   # 210mm × 297mm
    "Tabloid": (17280, 22320),   # 12in × 15.5in  (ANSI B)
}
MARGIN_NORMAL = 720   # 0.5 inch in twips
MARGIN_NARROW = 540   # 0.375 inch

def _set_cell_bg(cell, hex_color):
    """צבע רקע לתא"""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    # הסר shd קיים אם יש
    for old in tcPr.findall(qn('w:shd')):
        tcPr.remove(old)
    tcPr.append(shd)

def _set_cell_borders(cell, color="000000", size=4):
    """גבול דק לכל צדדי תא"""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    for old in tcPr.findall(qn('w:tcBorders')):
        tcPr.remove(old)
    tcBorders = OxmlElement('w:tcBorders')
    for side in ('top', 'left', 'bottom', 'right'):
        el = OxmlElement(f'w:{side}')
        el.set(qn('w:val'), 'single')
        el.set(qn('w:sz'), str(size))
        el.set(qn('w:space'), '0')
        el.set(qn('w:color'), color)
        tcBorders.append(el)
    tcPr.append(tcBorders)

def _cell_text(cell, text, bold=False, size=8, color=None, align="center",
               rtl=False, bg=None, valign="center"):
    """כתוב טקסט בתא עם עיצוב"""
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    if bg:
        _set_cell_bg(cell, bg)
    _set_cell_borders(cell)

    # נקה תוכן קיים
    for p in cell.paragraphs:
        for r in p.runs:
            r.text = ""

    p = cell.paragraphs[0]
    p.clear()
    if align == "center":
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    elif align == "right":
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    else:
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT

    # RTL
    if rtl:
        pPr = p._p.get_or_add_pPr()
        bidi = OxmlElement('w:bidi')
        pPr.append(bidi)

    run = p.add_run(str(text) if text is not None else "")
    run.bold = bold
    run.font.name = "Arial"
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = RGBColor.from_string(color)

def _set_section_props(section, page_size, landscape):
    """הגדר גודל דף וכיוון"""
    w_twips, h_twips = PAGE_SIZES[page_size]
    if landscape:
        section.page_width  = Twips(h_twips)
        section.page_height = Twips(w_twips)
        section.orientation = WD_ORIENT.LANDSCAPE
    else:
        section.page_width  = Twips(w_twips)
        section.page_height = Twips(h_twips)
        section.orientation = WD_ORIENT.PORTRAIT

    margin = MARGIN_NARROW
    section.top_margin    = Twips(margin)
    section.bottom_margin = Twips(margin)
    section.left_margin   = Twips(margin)
    section.right_margin  = Twips(margin)

def _content_width_twips(page_size, landscape):
    """רוחב תוכן שמיש בטוויפס"""
    w_twips, h_twips = PAGE_SIZES[page_size]
    if landscape:
        return h_twips - 2 * MARGIN_NARROW
    else:
        return w_twips - 2 * MARGIN_NARROW

def _add_section_break(doc, page_size, landscape):
    """הוסף מעבר דף + section חדש עם הגדרות"""
    new_section = doc.add_section()
    _set_section_props(new_section, page_size, landscape)
    return new_section

def _add_title_paragraph(doc, title, part_str, rtl=True):
    """כותרת טבלה מעל כל דף"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT if rtl else WD_ALIGN_PARAGRAPH.LEFT
    pPr = p._p.get_or_add_pPr()
    if rtl:
        bidi = OxmlElement('w:bidi')
        pPr.append(bidi)
    sp = p.paragraph_format
    sp.space_before = Pt(0)
    sp.space_after  = Pt(3)

    run = p.add_run(f"{title}   {part_str}")
    run.bold = True
    run.font.name = "Arial"
    run.font.size = Pt(11)
    return p

def _add_legend(doc, has_yellow, has_orange):
    """מקרא בתחתית טבלה"""
    if not has_yellow and not has_orange:
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    pPr = p._p.get_or_add_pPr()
    bidi = OxmlElement('w:bidi')
    pPr.append(bidi)
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after  = Pt(0)

    if has_yellow:
        run = p.add_run("■ ")
        run.font.name = "Arial"
        run.font.size = Pt(8)
        run.font.color.rgb = RGBColor.from_string(COLOR_YELLOW)
        run2 = p.add_run("בצהוב - חריגה מערך הסף VSL    ")
        run2.font.name = "Arial"
        run2.font.size = Pt(8)

    if has_orange:
        run3 = p.add_run("■ ")
        run3.font.name = "Arial"
        run3.font.size = Pt(8)
        run3.font.color.rgb = RGBColor.from_string(COLOR_ORANGE)
        run4 = p.add_run("בכתום - חריגה מערך הסף TIER 1")
        run4.font.name = "Arial"
        run4.font.size = Pt(8)

def _twips_to_emu(t):
    return int(t * 914400 / 1440)

def _set_table_width(table, width_twips):
    """הגדר רוחב טבלה"""
    tbl = table._tbl
    tblPr = tbl.find(qn('w:tblPr'))
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        tbl.insert(0, tblPr)
    for old in tblPr.findall(qn('w:tblW')):
        tblPr.remove(old)
    tblW = OxmlElement('w:tblW')
    tblW.set(qn('w:w'), str(width_twips))
    tblW.set(qn('w:type'), 'dxa')
    tblPr.append(tblW)

def _set_col_width(cell, width_twips):
    """רוחב תא"""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    for old in tcPr.findall(qn('w:tcW')):
        tcPr.remove(old)
    tcW = OxmlElement('w:tcW')
    tcW.set(qn('w:w'), str(width_twips))
    tcW.set(qn('w:type'), 'dxa')
    tcPr.append(tcW)

# ═══════════════════════════════════════════════════════════════════════════════
#  פונקציות בניית טבלאות לפי סוג
# ═══════════════════════════════════════════════════════════════════════════════

def _build_tph_table_data(df, thresh_dict, t1col, t1lbl):
    """
    מחזיר: headers (list of list), rows (list of dict with 'values' and 'colors')
    """
    import re

    def norm(s):
        s = "" if s is None else str(s).strip().lower()
        return re.sub(r"\s+", " ", s.replace("\xa0", " "))

    def sort_key(sid):
        m = re.match(r"S-?(\d+)", str(sid), re.I)
        return int(m.group(1)) if m else 9999

    def to_float(v):
        s = str(v).strip() if v is not None else ""
        try: return float(s.lstrip("<>").strip())
        except: return None

    def is_dro(c):
        if "(dro)" in c or "- dro" in c or c.strip()=="dro": return True
        if "(oro)" in c or "- oro" in c: return False
        if "c10" in c and "c28" in c and "c40" not in c: return True
        return False

    def is_oro(c):
        if "(oro)" in c or "- oro" in c or c.strip()=="oro": return True
        if "(dro)" in c or "- dro" in c: return False
        if "c24" in c and "c40" in c: return True
        if "c28" in c and "c40" in c: return True
        return False

    def is_total(c):
        if any(x in c for x in ["(dro)","(oro)","- dro","- oro"]): return False
        if "c10" in c and "c40" in c: return True
        if "total" in c and ("tph" in c or "hydrocarbon" in c): return True
        return False

    def get_thresh_local(compound, thresh_dict, t1col):
        from word_export import _match_thresh_simple
        t = _match_thresh_simple(compound, thresh_dict)
        return t.get("VSL"), t.get(t1col), t.get("cas", "-")

    vsl_d,t1_d,_ = get_thresh_local("C10 - C28 Fraction (DRO)", thresh_dict, t1col)
    vsl_o,t1_o,_ = get_thresh_local("C24 - C40 Fraction (ORO)", thresh_dict, t1col)
    vsl_t,t1_t,_ = get_thresh_local("TPH - DRO + ORO (Tier 1)", thresh_dict, t1col)
    vv = [v for v in [vsl_d,vsl_o,vsl_t] if v]
    tt = [v for v in [t1_d,t1_o,t1_t] if v]
    vsl_tot = min(float(v) for v in vv) if vv else 350
    t1_tot  = min(float(v) for v in tt) if tt else 350

    headers = [
        ["שם קידוח", "עומק", "TPH DRO", "TPH ORO", "Total TPH"],
        ["", "יחידות", "mg/kg", "mg/kg", "mg/kg"],
        ["", "CAS", "C10-C40", "C10-C40", "C10-C40"],
        ["", "VSL", str(vsl_tot), str(vsl_tot), str(vsl_tot)],
        ["", t1lbl.replace("\n", " "), str(t1_tot), str(t1_tot), str(t1_tot)],
    ]

    pivoted = {}
    for _, r in df.iterrows():
        k = (r["sample_id"], r["depth"])
        if k not in pivoted:
            pivoted[k] = {"DRO":"","ORO":"","TOT":"","DRO_f":None,"ORO_f":None,"DRO_lor":None,"ORO_lor":None}
        c = r["compound_lower"]
        if is_dro(c) and not pivoted[k]["DRO"]:
            pivoted[k]["DRO"] = r["result_str"]
            pivoted[k]["DRO_f"] = r["result"]
            pivoted[k]["DRO_lor"] = r.get("lor_val")
        elif is_oro(c) and not pivoted[k]["ORO"]:
            pivoted[k]["ORO"] = r["result_str"]
            pivoted[k]["ORO_f"] = r["result"]
            pivoted[k]["ORO_lor"] = r.get("lor_val")
        elif is_total(c) and not pivoted[k]["TOT"]:
            pivoted[k]["TOT"] = r["result_str"]

    def check_exceed(val_str, vsl, tier1):
        if not val_str or str(val_str).strip().startswith("<"): return None
        f = to_float(val_str)
        if f is None: return None
        try:
            t1f = float(tier1) if tier1 is not None else None
            vf  = float(vsl)   if vsl   is not None else None
            if t1f and t1f > 0 and f > t1f: return "tier1"
            if vf  and vf  > 0 and f > vf:  return "vsl"
        except: pass
        return None

    rows_out = []
    prev_sid = None
    for (sid,depth_val),v in sorted(pivoted.items(), key=lambda x:(sort_key(x[0][0]), x[0][1] or 0)):
        if v["TOT"]: total_s = v["TOT"]
        else:
            dro_lor = v["DRO"] and str(v["DRO"]).startswith("<")
            oro_lor = v["ORO"] and str(v["ORO"]).startswith("<")
            dro_empty = not v["DRO"]; oro_empty = not v["ORO"]
            dro_num = (v["DRO_lor"] if dro_lor and v["DRO_lor"] is not None else (v["DRO_f"] or 0))
            oro_num = (v["ORO_lor"] if oro_lor and v["ORO_lor"] is not None else (v["ORO_f"] or 0))
            total_f = dro_num + oro_num
            if (dro_lor or dro_empty) and (oro_lor or oro_empty) and not (dro_empty and oro_empty):
                total_s = f"<{total_f:.0f}"
            else:
                total_s = f"{total_f:.0f}"

        hl_dro   = check_exceed(v["DRO"],   vsl_tot, t1_tot)
        hl_oro   = check_exceed(v["ORO"],   vsl_tot, t1_tot)
        hl_total = check_exceed(total_s,    vsl_tot, t1_tot)

        sid_display = sid if sid != prev_sid else ""
        prev_sid = sid
        rows_out.append({
            "values": [sid_display, str(depth_val) if depth_val else "", v["DRO"], v["ORO"], total_s],
            "colors": [None, None, hl_dro, hl_oro, hl_total],
        })

    return headers, rows_out


def _build_metals_table_data(df, thresh_dict, t1col, t1lbl):
    import re
    METAL_MAP = {
        "aluminium":"Al","aluminum":"Al","antimony":"Sb","arsenic":"As",
        "barium":"Ba","beryllium":"Be","bismuth":"Bi","boron":"B",
        "cadmium":"Cd","calcium":"Ca","chromium":"Cr","cobalt":"Co",
        "copper":"Cu","iron":"Fe","lead":"Pb","lithium":"Li",
        "magnesium":"Mg","manganese":"Mn","mercury":"Hg","nickel":"Ni",
        "potassium":"K","selenium":"Se","silver":"Ag","sodium":"Na",
        "vanadium":"V","zinc":"Zn","molybdenum":"Mo","tin":"Sn",
        "titanium":"Ti","strontium":"Sr","thallium":"Tl",
        "phosphorus":"P","sulphur":"S","silicon":"Si",
    }
    METALS_ORDER = ["Al","Sb","As","Ba","Be","Bi","B","Cd","Ca","Cr","Co","Cu","Fe",
                    "Pb","Li","Mg","Mn","Hg","Ni","K","Se","Ag","Na","V","Zn"]
    THRESH_METAL_MAP = {
        "aluminum":"Al","antimony (metallic)":"Sb","antimony":"Sb",
        "arsenic, inorganic":"As","arsenic":"As","barium":"Ba",
        "beryllium and compounds":"Be","beryllium":"Be",
        "boron and borates only":"B","boron":"B",
        "cadmium (water) source: water and air":"Cd","cadmium":"Cd",
        "calcium":"Ca","chromium, total":"Cr","chromium":"Cr","cobalt":"Co",
        "copper":"Cu","iron":"Fe","lead and compounds":"Pb","lead":"Pb",
        "lithium":"Li","magnesium":"Mg","manganese (non-diet)":"Mn","manganese":"Mn",
        "mercuric chloride (and other mercury salts)":"Hg","mercury":"Hg",
        "nickel soluble salts":"Ni","nickel":"Ni","potassium":"K",
        "selenium":"Se","silver":"Ag","sodium":"Na",
        "vanadium and compounds":"V","vanadium":"V",
        "zinc and compounds":"Zn","zinc":"Zn","molybdenum":"Mo","tin":"Sn",
        "titanium":"Ti","strontium":"Sr","thallium":"Tl",
        "phosphorus":"P","sulphur":"S","silicon":"Si",
    }

    def norm(s):
        s = "" if s is None else str(s).strip().lower()
        return re.sub(r"\s+", " ", s.replace("\xa0", " "))

    def sort_key(sid):
        m = re.match(r"S-?(\d+)", str(sid), re.I)
        return int(m.group(1)) if m else 9999

    def to_float(v):
        s = str(v).strip() if v is not None else ""
        try: return float(s.lstrip("<>").strip())
        except: return None

    df = df.copy()
    df["sym"] = df["compound_lower"].map(METAL_MAP)
    df = df[df["sym"].notna()]
    if df.empty:
        return [["אין נתוני מתכות"]], []

    present = set(df["sym"].unique())
    metals = [m for m in METALS_ORDER if m in present] + sorted(present - set(METALS_ORDER))

    # build metals thresh
    mt = {}
    for key, v in thresh_dict.items():
        if key == "__CANONICAL_MAP_INTERNAL__": continue
        sym = THRESH_METAL_MAP.get(key)
        if sym and sym not in mt:
            mt[sym] = {"vsl": v.get("VSL"), "tier1": v.get(t1col), "cas": v.get("cas","-")}

    headers = [
        ["שם קידוח", "עומק"] + metals,
        ["", "יחידות"] + ["mg/kg"] * len(metals),
        ["", "CAS"] + [mt.get(m,{}).get("cas","-") for m in metals],
        ["", "VSL"] + [str(mt.get(m,{}).get("vsl","-")) for m in metals],
        ["", t1lbl.replace("\n"," ")] + [str(mt.get(m,{}).get("tier1","-")) for m in metals],
    ]

    pt = df.pivot_table(index=["sample_id","depth"], columns="sym", values="result_str", aggfunc="first")
    pt = pt.reindex(sorted(pt.index, key=lambda x:(sort_key(x[0]), x[1] or 0)))

    def check_exceed(val_str, vsl, tier1):
        if not val_str or str(val_str).strip().startswith("<"): return None
        f = to_float(val_str)
        if f is None: return None
        try:
            t1f = float(tier1) if tier1 is not None else None
            vf  = float(vsl)   if vsl   is not None else None
            if t1f and t1f > 0 and f > t1f: return "tier1"
            if vf  and vf  > 0 and f > vf:  return "vsl"
        except: pass
        return None

    rows_out = []
    prev_sid = None
    for (sid, depth_val), row_data in pt.iterrows():
        sid_display = sid if sid != prev_sid else ""
        prev_sid = sid
        values = [sid_display, str(depth_val) if depth_val else ""]
        colors = [None, None]
        for sym in metals:
            val = row_data.get(sym, "") or ""
            val = "" if str(val) == "nan" else str(val)
            hl = check_exceed(val, mt.get(sym,{}).get("vsl"), mt.get(sym,{}).get("tier1"))
            values.append(val)
            colors.append(hl)
        rows_out.append({"values": values, "colors": colors})

    return headers, rows_out


def _build_generic_table_data(df, thresh_dict, t1col, t1lbl, compound_order_key="compound"):
    """גנרי ל-VOC/SVOC ו-PFAS - מחזיר נתונים גולמיים"""
    import re

    def sort_key(sid):
        m = re.match(r"S-?(\d+)", str(sid), re.I)
        return int(m.group(1)) if m else 9999

    def to_float(v):
        s = str(v).strip() if v is not None else ""
        try: return float(s.lstrip("<>").strip())
        except: return None

    def check_exceed(val_str, vsl, tier1):
        if not val_str or str(val_str).strip().startswith("<"): return None
        f = to_float(val_str)
        if f is None: return None
        try:
            t1f = float(tier1) if tier1 is not None else None
            vf  = float(vsl)   if vsl   is not None else None
            if t1f and t1f > 0 and f > t1f: return "tier1"
            if vf  and vf  > 0 and f > vf:  return "vsl"
        except: pass
        return None

    pairs = sorted(df[["sample_id","depth"]].drop_duplicates().values.tolist(),
                   key=lambda x: (sort_key(x[0]), -(x[1] or 0)))

    compounds = df[compound_order_key].unique().tolist()

    headers = [
        ["שם התרכובת", "CAS", "VSL", t1lbl.replace("\n"," "), "יחידות"] + [f"{sid}\n{d}" for sid,d in pairs],
    ]

    als_data = {}
    for _, r in df.iterrows():
        k = str(r[compound_order_key]).strip()
        als_data.setdefault(k, {})[(r["sample_id"], r["depth"])] = r["result_str"]

    rows_out = []
    for cmp in compounds:
        t = _match_thresh_simple(cmp, thresh_dict)
        vsl   = t.get("VSL")
        tier1 = t.get(t1col)
        cas   = t.get("cas", "-")
        unit  = df[df[compound_order_key]==cmp]["unit"].iloc[0] if len(df[df[compound_order_key]==cmp]) > 0 else "mg/kg"
        values = [cmp, str(cas), str(vsl) if vsl else "-", str(tier1) if tier1 else "-", str(unit)]
        colors = [None, None, None, None, None]
        for sid, depth_val in pairs:
            rs = als_data.get(cmp, {}).get((sid, depth_val), "")
            hl = check_exceed(rs, vsl, tier1)
            values.append(str(rs) if rs else "")
            colors.append(hl)
        rows_out.append({"values": values, "colors": colors})

    return headers, rows_out


def _match_thresh_simple(compound, thresh_dict):
    """חיפוש פשוט במילון ערכי סף"""
    import re
    def norm(s):
        s = "" if s is None else str(s).strip().lower()
        return re.sub(r"\s+", " ", s.replace("\xa0", " "))
    key = norm(compound)
    for k in (key, key.replace(".",","), key.replace(",",".")):
        if k in thresh_dict and k != "__CANONICAL_MAP_INTERNAL__":
            return thresh_dict[k]
    return {}


# ═══════════════════════════════════════════════════════════════════════════════
#  פונקציה ראשית: כתוב טבלה ל-document
# ═══════════════════════════════════════════════════════════════════════════════

def _add_table_to_doc(doc, headers, data_rows, title, page_size, landscape,
                      is_first_section=False):
    """
    הוסף טבלה אחת (או חלק ממנה) למסמך.
    מחלק אוטומטית לדפים לפי מספר שורות שמתאימות.
    """
    if not data_rows:
        return

    n_cols = len(headers[0])
    content_w = _content_width_twips(page_size, landscape)

    # חישוב רוחב עמודות: עמודה ראשונה רחבה יותר
    if n_cols <= 5:
        col_w_first = int(content_w * 0.25)
        col_w_rest  = int((content_w - col_w_first) / max(n_cols - 1, 1))
        col_widths = [col_w_first] + [col_w_rest] * (n_cols - 1)
    else:
        col_w_first = int(content_w * 0.18)
        col_w_second = int(content_w * 0.08)
        remaining = content_w - col_w_first - col_w_second
        col_w_rest = int(remaining / max(n_cols - 2, 1))
        col_widths = [col_w_first, col_w_second] + [col_w_rest] * (n_cols - 2)

    # וודא שהסכום שווה לתוכן
    total = sum(col_widths)
    if total != content_w:
        col_widths[-1] += (content_w - total)

    # כמה שורות נתונים מתאימות לדף (הערכה)
    PAGE_H_TWIPS = PAGE_SIZES[page_size][1] if not landscape else PAGE_SIZES[page_size][0]
    available_h = PAGE_H_TWIPS - 2 * MARGIN_NARROW
    HEADER_ROWS_H = 600   # גובה כותרת טבלה בטוויפס
    TITLE_H = 400
    LEGEND_H = 300
    DATA_ROW_H = 280      # גובה שורת נתונים
    n_header_rows = len(headers)

    rows_per_page = max(5, int(
        (available_h - TITLE_H - LEGEND_H - n_header_rows * HEADER_ROWS_H) / DATA_ROW_H
    ))

    # פצל לחלקים
    chunks = []
    for i in range(0, len(data_rows), rows_per_page):
        chunks.append(data_rows[i:i+rows_per_page])

    total_parts = len(chunks)

    for part_idx, chunk in enumerate(chunks):
        part_num = part_idx + 1
        part_str = f"(חלק {part_num} מתוך {total_parts})" if total_parts > 1 else ""

        # section break (לא לחלק הראשון בראשון)
        if not (is_first_section and part_idx == 0):
            _add_section_break(doc, page_size, landscape)

        # כותרת
        _add_title_paragraph(doc, title, part_str, rtl=True)

        # בנה טבלה
        n_rows_total = n_header_rows + len(chunk)
        table = doc.add_table(rows=n_rows_total, cols=n_cols)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        _set_table_width(table, content_w)

        # כתוב header rows
        for hi, hrow in enumerate(headers):
            for ci, val in enumerate(hrow):
                cell = table.cell(hi, ci)
                _set_col_width(cell, col_widths[ci])
                is_first_hdr_row = (hi == 0)
                bg = COLOR_HEADER2 if is_first_hdr_row else COLOR_HEADER
                _cell_text(cell, val, bold=True, size=8, bg=bg, rtl=True)

        # כתוב data rows
        has_yellow = False
        has_orange = False
        for ri, row_data in enumerate(chunk):
            tr_idx = n_header_rows + ri
            for ci, (val, color) in enumerate(zip(row_data["values"], row_data["colors"])):
                cell = table.cell(tr_idx, ci)
                _set_col_width(cell, col_widths[ci])
                bg = None
                if color == "vsl":
                    bg = COLOR_YELLOW
                    has_yellow = True
                elif color == "tier1":
                    bg = COLOR_ORANGE
                    has_orange = True
                left_align = (ci == 0)
                _cell_text(cell, val, bold=bool(color), size=8, bg=bg or COLOR_WHITE,
                           align="right" if left_align else "center", rtl=True)

        # מקרא
        _add_legend(doc, has_yellow, has_orange)


# ═══════════════════════════════════════════════════════════════════════════════
#  פונקציה ציבורית: בנה קובץ Word
# ═══════════════════════════════════════════════════════════════════════════════

def build_word_report(table_configs, thresh_dict, t1col, t1lbl):
    """
    table_configs: list of dicts, each:
      {
        "type": "TPH" | "Metals" | "VOC+SVOC" | "PFAS",
        "df": DataFrame,
        "title": str,
        "page_size": "A4" | "Tabloid",
        "landscape": bool,
      }
    מחזיר: bytes של קובץ docx
    """
    doc = Document()

    # הסר section ראשוני ריק
    for section in doc.sections:
        section.page_width  = Twips(PAGE_SIZES["A4"][0])
        section.page_height = Twips(PAGE_SIZES["A4"][1])
        section.top_margin    = Twips(MARGIN_NARROW)
        section.bottom_margin = Twips(MARGIN_NARROW)
        section.left_margin   = Twips(MARGIN_NARROW)
        section.right_margin  = Twips(MARGIN_NARROW)

    # הסר פסקה ריקה ראשונית
    for p in doc.paragraphs:
        p._element.getparent().remove(p._element)

    is_first = True

    for cfg in table_configs:
        df        = cfg["df"]
        ttype     = cfg["type"]
        title     = cfg["title"]
        page_size = cfg.get("page_size", "A4")
        landscape = cfg.get("landscape", False)

        if df is None or df.empty:
            continue

        # בנה נתוני טבלה
        if ttype == "TPH":
            headers, data_rows = _build_tph_table_data(df, thresh_dict, t1col, t1lbl)
        elif ttype == "Metals":
            headers, data_rows = _build_metals_table_data(df, thresh_dict, t1col, t1lbl)
        elif ttype in ("VOC+SVOC", "PFAS"):
            headers, data_rows = _build_generic_table_data(df, thresh_dict, t1col, t1lbl)
        else:
            headers, data_rows = _build_generic_table_data(df, thresh_dict, t1col, t1lbl)

        if not data_rows:
            continue

        # הגדר section ראשון
        if is_first:
            section = doc.sections[0]
            _set_section_props(section, page_size, landscape)
            is_first = False

        _add_table_to_doc(doc, headers, data_rows, title, page_size, landscape,
                          is_first_section=(len(doc.element.body) <= 2))

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.getvalue()

tab_excel, tab_word = st.tabs(["📊 יצוא Excel", "📄 יצוא Word"])

# ══════════════════════════════════════════════════════════════════════════════
with tab_excel:
    st.sidebar.header("⚙️ הגדרות ערכי סף")
    st.sidebar.markdown("🟡 חריגה מ-VSL &nbsp;&nbsp;&nbsp; 🟠 חריגה מ-TIER 1")
    st.sidebar.markdown("---")
    land_use = st.sidebar.selectbox("Land Use", ["Industrial", "Residential"], index=0)
    aquifer  = st.sidebar.selectbox("Aquifer Sensitivity", ["A-1, A, B", "B-1 or C"], index=0)
    depth_opts = ["Not Applicable"] if "b-1" in aquifer.lower() else ["0 - 6 m", ">6 m"]
    depth    = st.sidebar.selectbox("Depth to Groundwater", depth_opts, index=0)
    t1col    = get_tier1_col(land_use, aquifer, depth)
    t1lbl    = tier1_label(land_use, aquifer, depth)
    st.sidebar.info(f"TIER 1: **{land_use}** | {aquifer} | {depth}")

    thresh_file = st.sidebar.file_uploader("📂 קובץ ערכי סף (Excel)", type=["xlsx","xls"], key="thresh")
    thresh_dict = {}
    if thresh_file:
        try:
            thresh_dict = load_threshold_file(thresh_file.read())
            st.sidebar.success(f"✅ נטענו {len(thresh_dict)-1} תרכובות")
        except Exception as e:
            st.sidebar.error(f"❌ שגיאה בטעינת קובץ סף: {e}")

    uploaded_files = st.file_uploader(
        "העלה קבצי ALS (Excel)", type=["xlsx","xls"],
        accept_multiple_files=True, key="als_files"
    )

    if uploaded_files:
        all_dfs = {"TPH": [], "Metals": [], "VOC": [], "PFAS": []}
        errors = []
        for uf in uploaded_files:
            try:
                result = parse_als_file(uf.read(), uf.name)
                for dtype, df in result.items():
                    if dtype in all_dfs and df is not None and len(df) > 0:
                        all_dfs[dtype].append(df)
            except Exception as e:
                errors.append(f"{uf.name}: {e}")

        if errors:
            for err in errors:
                st.warning(f"⚠️ {err}")

        import pandas as pd
        merged = {}
        for dtype, dfs in all_dfs.items():
            if dfs:
                merged[dtype] = pd.concat(dfs, ignore_index=True)

        if merged:
            from openpyxl import Workbook
            wb = Workbook()
            wb.remove(wb.active)

            def make_single_wb(write_fn, df, sname):
                wb2 = Workbook(); wb2.remove(wb2.active)
                write_fn(wb2.create_sheet(sname), df, thresh_dict, t1col, t1lbl)
                b = io.BytesIO(); wb2.save(b); b.seek(0)
                return b

            if "TPH" in merged:
                ws = wb.create_sheet("TPH")
                write_tph_sheet(ws, merged["TPH"], thresh_dict, t1col, t1lbl)
            if "Metals" in merged:
                ws = wb.create_sheet("Metals")
                write_metals_sheet(ws, merged["Metals"], thresh_dict, t1col, t1lbl)
            if "VOC" in merged:
                ws = wb.create_sheet("VOC_SVOC")
                write_voc_sheet(ws, merged["VOC"], thresh_dict, t1col, t1lbl)
            if "PFAS" in merged:
                ws = wb.create_sheet("PFAS")
                write_pfas_sheet(ws, merged["PFAS"], thresh_dict, t1col, t1lbl)

            buf = io.BytesIO(); wb.save(buf); buf.seek(0)
            st.success(f"✅ עובדו: {', '.join(merged.keys())}")
            st.download_button(
                "⬇️ הורד קובץ Excel מאוחד", data=buf.getvalue(),
                file_name="soil_report.xlsx",
                mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                use_container_width=True
            )

            cols = st.columns(4)
            singles = [
                ("TPH", write_tph_sheet, "TPH.xlsx"),
                ("Metals", write_metals_sheet, "Metals.xlsx"),
                ("VOC", write_voc_sheet, "VOC_SVOC.xlsx"),
                ("PFAS", write_pfas_sheet, "PFAS.xlsx"),
            ]
            for ci, (dtype, fn, fname) in enumerate(singles):
                if dtype in merged:
                    with cols[ci]:
                        b = make_single_wb(fn, merged[dtype], dtype)
                        st.download_button(
                            f"⬇️ {dtype}", data=b.getvalue(),
                            file_name=fname,
                            mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                            key=f"dl_{dtype.lower()}", use_container_width=True
                        )
        else:
            st.info("לא נמצאו נתונים בקבצים שהועלו")
    else:
        st.info("👆 העלה קבצי ALS לעיבוד")

# ══════════════════════════════════════════════════════════════════════════════
def build_tph_word(xl_file_bytes, table_num, page_size="A4", landscape=False):
    from openpyxl import load_workbook as lw
    import re as _re

    PAGE_DXA  = {"A4":(11906,16838), "Tabloid":(17280,22320)}
    MARGIN    = 720
    HDR_BLUE  = "B7D7F0"
    WHITE     = "FFFFFF"
    YELLOW    = "FFFF00"
    ORANGE    = "FFC000"
    N_HDR     = 5

    pw, ph = PAGE_DXA[page_size]
    page_w = ph if landscape else pw
    page_h = pw if landscape else ph
    content_w = page_w - 2*MARGIN

    N_COLS = 6
    ratios = [1480, 1060, 1060, 1680, 1680, 1680]
    rs = sum(ratios)
    col_ws = [int(content_w * r / rs) for r in ratios]
    col_ws[-1] += content_w - sum(col_ws)

    wb = lw(io.BytesIO(xl_file_bytes), data_only=True)
    ws = wb.active
    all_rows = list(ws.iter_rows(values_only=False))
    if len(all_rows) <= N_HDR:
        raise ValueError("הקובץ ריק")
    hdr_rows  = all_rows[:N_HDR]
    data_rows = all_rows[N_HDR:]

    def cv(r, c):
        v = hdr_rows[r][c].value
        return str(v).strip() if v is not None else ""

    vsl_vals   = [cv(3,2), cv(3,3), cv(3,4)]
    tier1_vals = [cv(4,2), cv(4,3), cv(4,4)]
    t1lbl      = cv(4,1) or t1lbl  # קרא label מהאקסל
    for i in range(3):
        if not vsl_vals[i]:
            vsl_vals[i] = next((v for v in vsl_vals if v), "350")
        if not tier1_vals[i]:
            tier1_vals[i] = next((v for v in tier1_vals if v), "350")

    # ── פורמט ערכים ──────────────────────────────────────────────────────────
    def fmt_val(raw):
        """פורמט ערך: פסיקים במספרים, מקף בשמות S1->S-1"""
        if raw is None: return ""
        s = str(raw).strip()
        if not s: return ""
        # מקף בשמות קידוח: אות(ות) + מספר → אות-מספר  e.g. S1->S-1, BH12->BH-12
        m = _re.match(r'^([A-Za-z]+)(\d+)$', s)
        if m:
            return f"{m.group(1)}-{m.group(2)}"
        # פסיקים במספרים שלמים / עשרוניים (לא ערכי "<X")
        if s.startswith('<'):
            return s
        try:
            f = float(s)
            if f == int(f) and abs(f) >= 1000:
                return f"{int(f):,}"
            elif abs(f) >= 1000:
                return f"{f:,.1f}"
            else:
                return s
        except:
            return s

    def is_eng_or_num(text):
        """אמת אם הטקסט הוא אנגלית/מספר (לא עברית)"""
        t = str(text).strip()
        if not t: return False
        has_heb = any('\u05d0' <= c <= '\u05ea' for c in t)
        return not has_heb

    def has_heb(text):
        """עברית → David. מספר טהור → David. אנגלית/מעורב → TNR."""
        t = str(text).strip()
        if not t: return False
        if any('\u05d0' <= c <= '\u05ea' for c in t): return True  # עברית
        if any(c.isalpha() for c in t): return False                  # יש אות לטינית
        return True  # רק מספרים/סמלים → David

    # ── קבץ שורות לפי קידוח ──────────────────────────────────────────────────
    drills = []
    current_sid = None
    current_group = []
    for xl_row in data_rows:
        sid = xl_row[0].value
        if sid is not None and str(sid).strip() != "" and sid != current_sid:
            if current_group:
                drills.append((current_sid, current_group))
            current_sid = sid
            current_group = [xl_row]
        else:
            current_group.append(xl_row)
    if current_group:
        drills.append((current_sid, current_group))

    # ── חישוב שורות לדף ──────────────────────────────────────────────────────
    ROW_H_TWIPS = 340
    HDR_H_TWIPS = 370
    SPACER_H    = 480   # רווח 1.5 שורה בין כותרת לטבלה
    TITLE_H     = 380
    LEGEND_H    = 360

    USABLE_H  = page_h - 2*MARGIN
    HDR_TOTAL = N_HDR * HDR_H_TWIPS
    avail     = USABLE_H - TITLE_H - SPACER_H - LEGEND_H - HDR_TOTAL
    rows_pp   = max(8, int(avail / ROW_H_TWIPS))

    # ── חלק לדפים - לא לפצל קידוח ────────────────────────────────────────────
    pages = []
    cur_page = []
    cur_count = 0
    for sid, drill_rows in drills:
        n = len(drill_rows)
        if cur_count > 0 and cur_count + n > rows_pp:
            pages.append(cur_page)
            cur_page = [(sid, drill_rows)]
            cur_count = n
        else:
            cur_page.append((sid, drill_rows))
            cur_count += n
    if cur_page:
        pages.append(cur_page)

    total_parts = len(pages)

    # ── helpers ────────────────────────────────────────────────────────────────
    def get_color(xl_cell):
        try:
            fill = xl_cell.fill
            if fill and fill.fill_type not in (None,"none",""):
                fg = fill.fgColor
                if fg.type == "rgb":
                    rgb = fg.rgb[-6:].upper()
                    if rgb not in ("000000","FFFFFF"): return rgb
        except: pass
        return None

    def set_page(section):
        if landscape:
            section.page_width  = Twips(page_h); section.page_height = Twips(page_w)
            section.orientation = WD_ORIENT.LANDSCAPE
        else:
            section.page_width  = Twips(page_w); section.page_height = Twips(page_h)
            section.orientation = WD_ORIENT.PORTRAIT
        for a in ('top_margin','bottom_margin','left_margin','right_margin'):
            setattr(section, a, Twips(MARGIN))

    def tbl_rtl(table):
        tbl = table._tbl
        tblPr = tbl.find(qn('w:tblPr'))
        if tblPr is None:
            tblPr = OxmlElement('w:tblPr'); tbl.insert(0, tblPr)
        for old in tblPr.findall(qn('w:tblW')): tblPr.remove(old)
        tblW = OxmlElement('w:tblW')
        tblW.set(qn('w:w'), str(content_w)); tblW.set(qn('w:type'),'dxa')
        tblPr.append(tblW)
        bv = OxmlElement('w:bidiVisual'); tblPr.append(bv)

    def set_cw(cell, w):
        tc = cell._tc; tcPr = tc.get_or_add_tcPr()
        for old in tcPr.findall(qn('w:tcW')): tcPr.remove(old)
        tcW = OxmlElement('w:tcW')
        tcW.set(qn('w:w'), str(w)); tcW.set(qn('w:type'),'dxa')
        tcPr.append(tcW)

    def set_bg(cell, hex_color):
        tc = cell._tc; tcPr = tc.get_or_add_tcPr()
        for old in tcPr.findall(qn('w:shd')): tcPr.remove(old)
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'),'clear'); shd.set(qn('w:color'),'auto')
        shd.set(qn('w:fill'), hex_color); tcPr.append(shd)

    def set_brd(cell):
        tc = cell._tc; tcPr = tc.get_or_add_tcPr()
        for old in tcPr.findall(qn('w:tcBorders')): tcPr.remove(old)
        tcB = OxmlElement('w:tcBorders')
        for s in ('top','left','bottom','right'):
            el = OxmlElement(f'w:{s}')
            el.set(qn('w:val'),'single'); el.set(qn('w:sz'),'4')
            el.set(qn('w:space'),'0'); el.set(qn('w:color'),'000000')
            tcB.append(el)
        tcPr.append(tcB)

    def set_vm(cell, restart=False):
        tc = cell._tc; tcPr = tc.get_or_add_tcPr()
        for old in tcPr.findall(qn('w:vMerge')): tcPr.remove(old)
        vm = OxmlElement('w:vMerge')
        if restart: vm.set(qn('w:val'),'restart')
        tcPr.append(vm)

    def set_row_h(row, h_twips):
        tr = row._tr
        trPr = tr.find(qn('w:trPr'))
        if trPr is None:
            trPr = OxmlElement('w:trPr'); tr.insert(0, trPr)
        for old in trPr.findall(qn('w:trHeight')): trPr.remove(old)
        trH = OxmlElement('w:trHeight')
        trH.set(qn('w:val'), str(h_twips))
        trH.set(qn('w:hRule'), 'exact')
        trPr.append(trH)

    def add_rtl_run(para, text, bold=False, size_heb=13, size_eng=11,
                    color_hex=None, underline=False):
        """XML-only run - תומך ב-newlines"""
        if not text: return
        from lxml import etree as _lxml_r
        _WR = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
        def _wr(tag): return f'{{{_WR}}}{tag}'

        use_heb = has_heb(text)
        fname   = "David" if use_heb else "Times New Roman"
        fsize   = str(size_heb * 2) if use_heb else str(size_eng * 2)

        lines = str(text).split('\n')

        def _make_run(txt):
            r_el   = _lxml_r.SubElement(para._p, _wr('r'))
            rPr_el = _lxml_r.SubElement(r_el,   _wr('rPr'))
            rF = _lxml_r.SubElement(rPr_el, _wr('rFonts'))
            for a in ('ascii','hAnsi','cs','eastAsia'): rF.set(_wr(a), fname)
            if bold:
                b  = _lxml_r.SubElement(rPr_el, _wr('b'));  b.set(_wr('val'), '1')
                bC = _lxml_r.SubElement(rPr_el, _wr('bCs')); bC.set(_wr('val'), '1')
            if underline:
                u = _lxml_r.SubElement(rPr_el, _wr('u')); u.set(_wr('val'), 'single')
            if color_hex:
                col = _lxml_r.SubElement(rPr_el, _wr('color')); col.set(_wr('val'), color_hex)
            sz   = _lxml_r.SubElement(rPr_el, _wr('sz'));   sz.set(_wr('val'), fsize)
            szCs = _lxml_r.SubElement(rPr_el, _wr('szCs')); szCs.set(_wr('val'), fsize)
            _lxml_r.SubElement(rPr_el, _wr('rtl'))
            t_el = _lxml_r.SubElement(r_el, _wr('t'))
            t_el.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
            t_el.text = txt

        def _make_br():
            br_r = _lxml_r.SubElement(para._p, _wr('r'))
            _lxml_r.SubElement(br_r, _wr('br'))

        for i, line in enumerate(lines):
            if i > 0:
                _make_br()
            _make_run(line)

    def make_rtl_para(doc, align='center'):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after  = Pt(0)
        pPr = p._p.get_or_add_pPr()
        pPr.append(OxmlElement('w:bidi'))
        jc = OxmlElement('w:jc'); jc.set(qn('w:val'), align); pPr.append(jc)
        return p

    def write_cell(cell, text, bold=False, bg=WHITE):
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        set_bg(cell, bg); set_brd(cell)
        p = cell.paragraphs[0]; p.clear()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after  = Pt(0)
        pPr = p._p.get_or_add_pPr()
        pPr.append(OxmlElement('w:bidi'))
        jc = OxmlElement('w:jc'); jc.set(qn('w:val'),'center'); pPr.append(jc)
        txt = fmt_val(text) if not any('\u05d0'<=c<='\u05ea' for c in str(text or '')) else str(text or '')
        if not txt: return
        add_rtl_run(p, txt, bold=bold, size_heb=13, size_eng=11)

    def build_header(table):
        for hi in range(N_HDR):
            set_row_h(table.rows[hi], HDR_H_TWIPS)
        write_cell(table.cell(0,0), "שם קידוח", bold=True, bg=HDR_BLUE)
        set_vm(table.cell(0,0), restart=True)
        write_cell(table.cell(0,1), "עומק", bold=True, bg=HDR_BLUE)
        set_vm(table.cell(0,1), restart=True)
        write_cell(table.cell(0,2), "אנליזה", bold=True, bg=HDR_BLUE)
        for ci,lbl in enumerate(["TPH DRO","TPH ORO","Total TPH"],3):
            write_cell(table.cell(0,ci), lbl, bold=True, bg=HDR_BLUE)
        sub = [("יחידות",["mg/kg","mg/kg","mg/kg"]),
               ("CAS",   ["C-10-C-40","C-10-C-40","C-10-C-40"]),
               ("VSL",   vsl_vals),
               (t1lbl,tier1_vals)]
        for hi,(lbl,vals) in enumerate(sub,1):
            write_cell(table.cell(hi,0),"",bg=HDR_BLUE); set_vm(table.cell(hi,0))
            c12=table.cell(hi,1); c12.merge(table.cell(hi,2))
            write_cell(c12,lbl,bold=True,bg=HDR_BLUE)
            for ci,v in enumerate(vals,3):
                write_cell(table.cell(hi,ci),v,bold=True,bg=HDR_BLUE)

    # ── בנה Document ─────────────────────────────────────────────────────────
    doc = Document()
    # הסר bold=false מסגנון ברירת מחדל
    from lxml import etree as _lxml_def
    _W_DEF = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
    styles = doc.element.find(f'{{{_W_DEF}}}styles')
    if styles is not None:
        for b_off in styles.findall(f'.//{{{_W_DEF}}}b[@{{{_W_DEF}}}val="0"]'):
            b_off.getparent().remove(b_off)
        for b_off in styles.findall(f'.//{{{_W_DEF}}}bCs[@{{{_W_DEF}}}val="0"]'):
            b_off.getparent().remove(b_off)
    for p in doc.paragraphs:
        p._element.getparent().remove(p._element)

    for part_idx, page_drills in enumerate(pages):
        pn  = part_idx + 1
        pts = f"(חלק {pn} מתוך {total_parts})" if total_parts > 1 else ""
        title = f"טבלה מספר {table_num} {pts} – תוצאות אנליזות TPH"

        section = doc.sections[0] if part_idx == 0 else doc.add_section()
        set_page(section)

        # כותרת - XML ידני מוחלט עם w: prefix
        from lxml import etree as _lxml
        _W = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
        def _w(tag): return f'{{{_W}}}{tag}'

        tp = doc.add_paragraph()
        tp.paragraph_format.space_before = Pt(0)
        tp.paragraph_format.space_after  = Pt(0)
        tp.paragraph_format.keep_with_next = True
        # נקה pPr קיים ובנה מחדש
        pPr_t = tp._p.get_or_add_pPr()
        # spacing
        sp = _lxml.SubElement(pPr_t, _w('spacing'))
        sp.set(_w('before'), '0'); sp.set(_w('after'), '0')
        # keepNext
        _lxml.SubElement(pPr_t, _w('keepNext'))
        # bidi paragraph
        _lxml.SubElement(pPr_t, _w('bidi'))
        # center
        jc_el = _lxml.SubElement(pPr_t, _w('jc'))
        jc_el.set(_w('val'), 'center')

        # פצל כותרת לפי אנגלית/עברית-מספרים וצור run לכל חלק
        import re as _re_title
        def _title_segments(txt):
            parts = _re_title.split(r'([A-Za-z]+)', txt)
            return [(p, bool(_re_title.match(r'^[A-Za-z]+$', p))) for p in parts if p]

        def _add_title_run(parent_p, text, is_eng):
            fname = 'Times New Roman' if is_eng else 'David'
            fsize = '22' if is_eng else '26'  # 11pt=22, 13pt=26 half-points
            r_el  = _lxml.SubElement(parent_p, _w('r'))
            rPr_el = _lxml.SubElement(r_el, _w('rPr'))
            rF = _lxml.SubElement(rPr_el, _w('rFonts'))
            for a in ('ascii','hAnsi','cs','eastAsia'): rF.set(_w(a), fname)
            b_el = _lxml.SubElement(rPr_el, _w('b')); b_el.set(_w('val'), '1')
            bCs_el = _lxml.SubElement(rPr_el, _w('bCs')); bCs_el.set(_w('val'), '1')
            u_el = _lxml.SubElement(rPr_el, _w('u')); u_el.set(_w('val'), 'single')
            sz_el = _lxml.SubElement(rPr_el, _w('sz')); sz_el.set(_w('val'), fsize)
            szCs_el = _lxml.SubElement(rPr_el, _w('szCs')); szCs_el.set(_w('val'), fsize)
            _lxml.SubElement(rPr_el, _w('rtl'))
            t_el = _lxml.SubElement(r_el, _w('t'))
            t_el.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
            t_el.text = text

        for seg, is_eng in _title_segments(title):
            _add_title_run(tp._p, seg, is_eng)

        # רווח קבוע 1.5 שורה בין כותרת לטבלה
        sp = doc.add_paragraph()
        sp.paragraph_format.space_before = Pt(0)
        sp.paragraph_format.space_after  = Pt(0)
        sp.paragraph_format.line_spacing = Pt(19.5)  # 1.5 שורות של 13pt
        sp.paragraph_format.keep_with_next = True

        # טבלה
        n_data = sum(len(dr) for _, dr in page_drills)
        table  = doc.add_table(rows=N_HDR + n_data, cols=N_COLS)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        tbl_rtl(table)
        for ri in range(N_HDR + n_data):
            for ci in range(N_COLS):
                set_cw(table.cell(ri,ci), col_ws[ci])
        build_header(table)

        # נתונים
        has_yel = False; has_org = False
        ri = 0
        for sid, drill_rows in page_drills:
            for sub_ri, xl_row in enumerate(drill_rows):
                row_idx = N_HDR + ri
                dep = xl_row[1].value
                sid_fmt = fmt_val(sid) if sid else ""
                if sub_ri == 0:
                    write_cell(table.cell(row_idx,0), sid_fmt, bold=True, bg=WHITE)
                    set_vm(table.cell(row_idx,0), restart=True)
                else:
                    write_cell(table.cell(row_idx,0), "", bg=WHITE)
                    set_vm(table.cell(row_idx,0))
                set_row_h(table.rows[row_idx], ROW_H_TWIPS)
                cd = table.cell(row_idx,1); cd.merge(table.cell(row_idx,2))
                dep_fmt = fmt_val(dep)
                write_cell(cd, dep_fmt, bg=WHITE)
                for ci in range(3,6):
                    xl_c = xl_row[ci] if ci < len(xl_row) else None
                    val  = xl_c.value if xl_c else ""
                    bg   = get_color(xl_c) if xl_c else None
                    if bg is None: bg = WHITE
                    is_exceed = bg in (YELLOW, ORANGE)
                    if bg == YELLOW: has_yel = True
                    if bg == ORANGE: has_org = True
                    write_cell(table.cell(row_idx,ci), str(val) if val is not None else "", bg=bg, bold=is_exceed)
                ri += 1

        # מקרא - ימין מוחלט: paragraph רגיל + RTL mark + right align
        from lxml import etree as _lxml2
        _W2 = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
        def _w2(tag): return f'{{{_W2}}}{tag}'

        lp = doc.add_paragraph()
        lp.paragraph_format.space_before = Pt(4)
        lp.paragraph_format.space_after  = Pt(0)
        # בנה pPr מחדש לגמרי - ללא bidi, עם right align בלבד
        # בפסקת LTR: jc=right = ויזואלית ימין (ללא פרדוקס bidi)
        lp_pPr = lp._p.get_or_add_pPr()
        # נקה כל jc קיים
        for old_jc in lp_pPr.findall(_w2('jc')): lp_pPr.remove(old_jc)
        for old_bi in lp_pPr.findall(_w2('bidi')): lp_pPr.remove(old_bi)
        jc_lp = _lxml2.SubElement(lp_pPr, _w2('jc'))
        jc_lp.set(_w2('val'), 'right')
        # ללא w:bidi - פסקת LTR עם right align = ויזואלית ימין

        def leg_word_run(para, word, hex_color):
            """מילה מודגשת עם צבע רקע מדויק"""
            r_el = _lxml2.SubElement(para._p, _w2('r'))
            rPr_el = _lxml2.SubElement(r_el, _w2('rPr'))
            rF = _lxml2.SubElement(rPr_el, _w2('rFonts'))
            for a in ('ascii','hAnsi','cs','eastAsia'): rF.set(_w2(a), 'David')
            b_el = _lxml2.SubElement(rPr_el, _w2('b'))
            b_el.set(_w2('val'), '1')
            bCs_el = _lxml2.SubElement(rPr_el, _w2('bCs'))
            bCs_el.set(_w2('val'), '1')
            sz = _lxml2.SubElement(rPr_el, _w2('sz')); sz.set(_w2('val'), '20')
            szC = _lxml2.SubElement(rPr_el, _w2('szCs')); szC.set(_w2('val'), '20')
            # צבע רקע מדויק עם shading
            shd = _lxml2.SubElement(rPr_el, _w2('shd'))
            shd.set(_w2('val'), 'clear')
            shd.set(_w2('color'), 'auto')
            shd.set(_w2('fill'), hex_color)
            _lxml2.SubElement(rPr_el, _w2('rtl'))
            t = _lxml2.SubElement(r_el, _w2('t'))
            t.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
            t.text = word

        def leg_plain_run(para, text):
            """טקסט רגיל"""
            r_el = _lxml2.SubElement(para._p, _w2('r'))
            rPr_el = _lxml2.SubElement(r_el, _w2('rPr'))
            rF = _lxml2.SubElement(rPr_el, _w2('rFonts'))
            for a in ('ascii','hAnsi','cs','eastAsia'): rF.set(_w2(a), 'David')
            sz = _lxml2.SubElement(rPr_el, _w2('sz')); sz.set(_w2('val'), '20')
            szC = _lxml2.SubElement(rPr_el, _w2('szCs')); szC.set(_w2('val'), '20')
            _lxml2.SubElement(rPr_el, _w2('rtl'))
            t = _lxml2.SubElement(r_el, _w2('t'))
            t.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
            t.text = text

        if has_yel:
            leg_word_run(lp, "בצהוב", "FFFF00")
            leg_plain_run(lp, " - חריגה מערך הסף VSL")
        if has_yel and has_org:
            leg_plain_run(lp, "     ")
        if has_org:
            leg_word_run(lp, "בכתום", "FFC000")
            leg_plain_run(lp, " - חריגה מערך הסף TIER 1")

    buf = io.BytesIO(); doc.save(buf); buf.seek(0)
    return buf.getvalue()


def build_metals_word(xl_file_bytes, table_num, page_size="A3", landscape=True, t1lbl="TIER 1"):
    """בונה דוח Word לטבלת מתכות - זהה לTPH עם התאמות"""
    import io, re
    from lxml import etree as _lxml
    from docx import Document
    from docx.shared import Pt, RGBColor, Twips, Inches, Emu
    from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
    from docx.enum.section import WD_ORIENT
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    import openpyxl

    YELLOW = "FFFF00"
    ORANGE = "FFC000"
    WHITE  = "FFFFFF"
    N_HDR  = 5   # שורות כותרת: שמות, יחידות, CAS, VSL, TIER1

    # ── קרא Excel ────────────────────────────────────────────────────────────
    wb = openpyxl.load_workbook(io.BytesIO(xl_file_bytes), data_only=True)
    ws = None
    for name in wb.sheetnames:
        if "metal" in name.lower():
            ws = wb[name]; break
    if ws is None:
        ws = wb.active
    if ws is None or ws.max_row < 2:
        raise ValueError("הקובץ ריק")

    xl_rows = list(ws.iter_rows())
    if not xl_rows:
        raise ValueError("הקובץ ריק")

    # סנן עמודות ריקות לגמרי
    n_raw = len(xl_rows[0])
    non_empty_cols = []
    for ci in range(n_raw):
        if any(xl_rows[ri][ci].value not in (None, "", "-") for ri in range(len(xl_rows))):
            non_empty_cols.append(ci)
    # בנה מחדש xl_rows רק עם עמודות לא ריקות
    xl_rows = [[row[ci] for ci in non_empty_cols] for row in xl_rows]

    n_cols = len(xl_rows[0])  # מספר עמודות כולל שם קידוח + עומק + מתכות

    # קרא t1lbl מהאקסל (שורה 4, עמודה 1 = תווית TIER 1)
    if len(xl_rows) > 4 and len(xl_rows[4]) > 1:
        xl_t1 = xl_rows[4][1].value
        if xl_t1 and str(xl_t1).strip():
            t1lbl = str(xl_t1).strip()

    def get_color(cell):
        try:
            fill = cell.fill
            if fill and fill.fill_type not in (None,"none"):
                fg = fill.fgColor
                if fg.type=="rgb" and fg.rgb not in ("00000000","FFFFFFFF","FF000000","FFFFFFFF"):
                    h = fg.rgb[-6:].upper()
                    if h == YELLOW: return YELLOW
                    if h == ORANGE: return ORANGE
            return None
        except: return None

    def fmt_val(raw):
        if raw is None: return ""
        s = str(raw).strip()
        if not s or s=="-": return s
        m = re.match(r'^([A-Za-z]+)(\d+)$', s)
        if m: return f"{m.group(1)}-{m.group(2)}"
        if s.startswith('<'): return s
        try:
            f = float(s)
            if f == int(f) and abs(f) >= 1000: return f"{int(f):,}"
            elif abs(f) >= 1000: return f"{f:,.1f}"
            else: return s
        except: return s

    def has_heb(text):
        t = str(text).strip()
        if not t: return False
        if any('\u05d0' <= c <= '\u05ea' for c in t): return True
        if any(c.isalpha() for c in t): return False
        return True

    _W = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
    def _w(tag): return f'{{{_W}}}{tag}'

    def set_page(section):
        if landscape:
            section.orientation = WD_ORIENT.LANDSCAPE
            section.page_width  = Inches(17)
            section.page_height = Inches(11)
        else:
            section.orientation = WD_ORIENT.PORTRAIT
            section.page_width  = Inches(8.5)
            section.page_height = Inches(11)
        section.left_margin   = Inches(0.4)
        section.right_margin  = Inches(0.4)
        section.top_margin    = Inches(0.4)
        section.bottom_margin = Inches(0.4)

    def set_bg(cell, hex_color):
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        for old in tcPr.findall(qn('w:shd')): tcPr.remove(old)
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'),   'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'),  hex_color)
        tcPr.append(shd)

    def set_brd(cell):
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        for old in tcPr.findall(qn('w:tcBorders')): tcPr.remove(old)
        borders = OxmlElement('w:tcBorders')
        for side in ('top','left','bottom','right'):
            b = OxmlElement(f'w:{side}')
            b.set(qn('w:val'),'single'); b.set(qn('w:sz'),'4')
            b.set(qn('w:space'),'0'); b.set(qn('w:color'),'000000')
            borders.append(b)
        tcPr.append(borders)

    def set_row_height(tr, h_twips):
        trPr = tr.find(qn('w:trPr'))
        if trPr is None:
            trPr = OxmlElement('w:trPr'); tr.insert(0, trPr)
        for old in trPr.findall(qn('w:trHeight')): trPr.remove(old)
        trH = OxmlElement('w:trHeight')
        trH.set(qn('w:val'), str(h_twips))
        trH.set(qn('w:hRule'), 'exact')
        trPr.append(trH)

    def add_rtl_run(para, text, bold=False, size_heb=9, size_eng=7,
                    color_hex=None, underline=False):
        """XML-only run - תומך ב-newlines"""
        if not str(text): return
        lines = str(text).split('\n')
        use_heb = has_heb(text)
        fname   = "David" if use_heb else "Times New Roman"
        fsize   = str(size_heb * 2) if use_heb else str(size_eng * 2)

        def _make_run(txt):
            r_el   = _lxml.SubElement(para._p, _w('r'))
            rPr_el = _lxml.SubElement(r_el,   _w('rPr'))
            rF = _lxml.SubElement(rPr_el, _w('rFonts'))
            for a in ('ascii','hAnsi','cs','eastAsia'): rF.set(_w(a), fname)
            if bold:
                b  = _lxml.SubElement(rPr_el, _w('b'));  b.set(_w('val'), '1')
                bC = _lxml.SubElement(rPr_el, _w('bCs')); bC.set(_w('val'), '1')
            if underline:
                u = _lxml.SubElement(rPr_el, _w('u')); u.set(_w('val'), 'single')
            if color_hex:
                col = _lxml.SubElement(rPr_el, _w('color')); col.set(_w('val'), color_hex)
            sz   = _lxml.SubElement(rPr_el, _w('sz'));   sz.set(_w('val'), fsize)
            szCs = _lxml.SubElement(rPr_el, _w('szCs')); szCs.set(_w('val'), fsize)
            _lxml.SubElement(rPr_el, _w('rtl'))
            t_el = _lxml.SubElement(r_el, _w('t'))
            t_el.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
            t_el.text = txt

        def _make_br():
            """שורה חדשה בתוך תא (w:br)"""
            br_r = _lxml.SubElement(para._p, _w('r'))
            _lxml.SubElement(br_r, _w('br'))

        for i, line in enumerate(lines):
            if i > 0:
                _make_br()
            _make_run(line)

    def write_cell(cell, text, bold=False, bg=WHITE, size_heb=9, size_eng=7):
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        set_bg(cell, bg); set_brd(cell)
        p = cell.paragraphs[0]; p.clear()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after  = Pt(0)
        pPr = p._p.get_or_add_pPr()
        pPr.append(OxmlElement('w:bidi'))
        jc = OxmlElement('w:jc'); jc.set(qn('w:val'),'center'); pPr.append(jc)
        raw = str(text or '')
        txt = raw if any('\u05d0'<=c<='\u05ea' for c in raw) else fmt_val(raw)
        if not txt: return
        add_rtl_run(p, txt, bold=bold, size_heb=size_heb, size_eng=size_eng)

    # ── קרא נתונים מ-Excel ───────────────────────────────────────────────────
    hdr_row = xl_rows[0]
    n_data_cols = len(hdr_row)

    # מצא שורות נתונים (אחרי 5 כותרות)
    data_xl_rows = xl_rows[N_HDR:]

    # קבץ לדפים - כמה שורות נכנסות בדף A3 landscape
    # A3 landscape: 11" height, margins 0.4" → usable ~10.2" = ~14688 twips
    ROW_H   = 280   # twips per data row
    HDR_H   = 320   # twips per header row
    TITLE_H = 400
    SPACER_H= 200
    LEGEND_H= 320
    usable  = 14688
    hdr_total = N_HDR * HDR_H
    rows_pp = int((usable - TITLE_H - SPACER_H - LEGEND_H - hdr_total) / ROW_H)
    rows_pp = max(10, rows_pp)

    # ── קבץ לפי קידוח ────────────────────────────────────────────────────────
    drills = []
    current_sid = None
    for xl_row in data_xl_rows:
        sid = xl_row[0].value
        if sid and str(sid).strip():
            current_sid = str(sid).strip()
        if current_sid is None: continue
        drills.append((current_sid, xl_row))

    # חלק לדפים ללא שבירת קידוח
    pages = []
    page  = []
    for item in drills:
        page.append(item)
        if len(page) >= rows_pp:
            pages.append(page); page = []
    if page: pages.append(page)
    if not pages: pages = [[]]

    total_parts = len(pages)
    title_base  = f"טבלה מספר {table_num}"

    # ── בנה מסמך ─────────────────────────────────────────────────────────────
    doc = Document()
    # הסר bold=false מסגנון ברירת מחדל
    _W_DEF = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
    styles = doc.element.find(f'{{{_W_DEF}}}styles')
    if styles is not None:
        for b_off in styles.findall(f'.//{{{_W_DEF}}}b[@{{{_W_DEF}}}val="0"]'):
            b_off.getparent().remove(b_off)
        for b_off in styles.findall(f'.//{{{_W_DEF}}}bCs[@{{{_W_DEF}}}val="0"]'):
            b_off.getparent().remove(b_off)
    for p in doc.paragraphs:
        p._element.getparent().remove(p._element)

    def _title_segments(txt):
        parts = re.split(r'([A-Za-z]+)', txt)
        return [(p, bool(re.match(r'^[A-Za-z]+$', p))) for p in parts if p]

    def _add_title_run(parent_p, text, is_eng):
        fname = 'Times New Roman' if is_eng else 'David'
        fsize = '22' if is_eng else '26'
        r_el   = _lxml.SubElement(parent_p, _w('r'))
        rPr_el = _lxml.SubElement(r_el,    _w('rPr'))
        rF = _lxml.SubElement(rPr_el, _w('rFonts'))
        for a in ('ascii','hAnsi','cs','eastAsia'): rF.set(_w(a), fname)
        b_el = _lxml.SubElement(rPr_el, _w('b')); b_el.set(_w('val'), '1')
        bCs_el = _lxml.SubElement(rPr_el, _w('bCs')); bCs_el.set(_w('val'), '1')
        u_el = _lxml.SubElement(rPr_el, _w('u')); u_el.set(_w('val'), 'single')
        sz_el = _lxml.SubElement(rPr_el, _w('sz')); sz_el.set(_w('val'), fsize)
        szCs_el = _lxml.SubElement(rPr_el, _w('szCs')); szCs_el.set(_w('val'), fsize)
        _lxml.SubElement(rPr_el, _w('rtl'))
        t_el = _lxml.SubElement(r_el, _w('t'))
        t_el.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
        t_el.text = text

    for part_idx, page_rows in enumerate(pages):
        pn    = part_idx + 1
        pts   = f"(חלק {pn} מתוך {total_parts})" if total_parts > 1 else ""
        title = f"{title_base} {pts} – תוצאות אנליזות מתכות".strip()

        section = doc.sections[0] if part_idx == 0 else doc.add_section()
        set_page(section)

        # ── כותרת ────────────────────────────────────────────────────────────
        tp = doc.add_paragraph()
        tp.paragraph_format.space_before = Pt(0)
        tp.paragraph_format.space_after  = Pt(0)
        tp.paragraph_format.keep_with_next = True
        pPr_t = tp._p.get_or_add_pPr()
        sp = _lxml.SubElement(pPr_t, _w('spacing'))
        sp.set(_w('before'), '0'); sp.set(_w('after'), '0')
        _lxml.SubElement(pPr_t, _w('keepNext'))
        _lxml.SubElement(pPr_t, _w('bidi'))
        jc_el = _lxml.SubElement(pPr_t, _w('jc')); jc_el.set(_w('val'), 'center')
        for seg, is_eng in _title_segments(title):
            _add_title_run(tp._p, seg, is_eng)

        # ── רווח ─────────────────────────────────────────────────────────────
        sp_p = doc.add_paragraph()
        sp_p.paragraph_format.space_before = Pt(0)
        sp_p.paragraph_format.space_after  = Pt(0)
        sp_p.paragraph_format.keep_with_next = True
        sp_p.paragraph_format.line_spacing = Pt(19.5)  # 1.5 שורות של 13pt

        # ── טבלה ─────────────────────────────────────────────────────────────
        n_rows_table = N_HDR + len(page_rows)
        table = doc.add_table(rows=n_rows_table, cols=n_data_cols)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.style = 'Table Grid'

        # רוחב עמודות
        total_w_emu = int((17 - 0.8) * 914400)  # usable width
        col0_w = int(total_w_emu * 0.05)
        col1_w = int(total_w_emu * 0.04)
        rest_w = (total_w_emu - col0_w - col1_w) // max(1, n_data_cols - 2)
        for ci in range(n_data_cols):
            w = col0_w if ci==0 else col1_w if ci==1 else rest_w
            for cell in table.columns[ci].cells:
                cell.width = Emu(w)

        # bidiVisual על הטבלה
        tblPr = table._tbl.find(qn('w:tblPr'))
        if tblPr is None:
            tblPr = OxmlElement('w:tblPr'); table._tbl.insert(0, tblPr)
        bv = OxmlElement('w:bidiVisual'); tblPr.append(bv)

        # ── 5 שורות כותרת ────────────────────────────────────────────────────
        HDR_BG = "B7D7F0"
        for hi in range(N_HDR):
            xl_hdr_row = xl_rows[hi]
            set_row_height(table.rows[hi]._tr, HDR_H)
            for ci in range(n_data_cols):
                xl_c = xl_hdr_row[ci] if ci < len(xl_hdr_row) else None
                val  = xl_c.value if xl_c else ""
                val_s = str(val).strip() if val is not None else ""
                # שורות 1-4 בעמודה 0: ריק (ימוזג)
                if ci == 0 and hi > 0:
                    val_s = ""
                # שורה 4 עמודה 1: החלף TIER 1 ב-t1lbl
                if hi == 4 and ci == 1:
                    val_s = t1lbl
                write_cell(table.cell(hi, ci), val_s, bold=True, bg=HDR_BG,
                           size_heb=9, size_eng=7)

        # מזג שורות 0-4 בעמודה 0 (שם קידוח) - vMerge
        for hi in range(N_HDR):
            tc   = table.cell(hi, 0)._tc
            tcPr = tc.get_or_add_tcPr()
            for old in tcPr.findall(qn('w:vMerge')): tcPr.remove(old)
            vm = OxmlElement('w:vMerge')
            if hi == 0:
                vm.set(qn('w:val'), 'restart')
            tcPr.append(vm)

        # מזג שורות 0-1 בעמודה 1 (עומק - ללא יחידות)
        for hi in range(2):
            tc   = table.cell(hi, 1)._tc
            tcPr = tc.get_or_add_tcPr()
            for old in tcPr.findall(qn('w:vMerge')): tcPr.remove(old)
            vm = OxmlElement('w:vMerge')
            if hi == 0:
                vm.set(qn('w:val'), 'restart')
            tcPr.append(vm)
        # תא עומק שורה 1 כבר ריק מה-Excel (אין יחידות לעומק)

        # ── שורות נתונים ─────────────────────────────────────────────────────
        has_yel = False; has_org = False
        # בנה מיפוי: כמה שורות יש לכל קידוח (לצורך vMerge)
        drill_row_counts = {}
        for sid, _ in page_rows:
            drill_row_counts[sid] = drill_row_counts.get(sid, 0) + 1
        drill_start = {}  # sid → row_idx ראשון
        drill_seen  = {}  # sid → כמה שורות כבר נכתבו

        for ri, (sid, xl_row) in enumerate(page_rows):
            row_idx = N_HDR + ri
            set_row_height(table.rows[row_idx]._tr, ROW_H)

            # vMerge על שם קידוח אם יש כמה שורות לאותו קידוח
            n_rows_for_drill = drill_row_counts.get(sid, 1)
            seen = drill_seen.get(sid, 0)
            drill_seen[sid] = seen + 1

            for ci in range(n_data_cols):
                xl_c = xl_row[ci] if ci < len(xl_row) else None
                val  = xl_c.value if xl_c else ""
                bg   = get_color(xl_c) if xl_c else None
                if bg is None: bg = WHITE
                is_exceed = bg in (YELLOW, ORANGE)
                if bg == YELLOW: has_yel = True
                if bg == ORANGE: has_org = True
                # col0 = שם קידוח: הצג רק בשורה הראשונה, מזג אם יש כמה
                if ci == 0:
                    cell_val = str(val) if val is not None else ""
                    write_cell(table.cell(row_idx, ci), cell_val,
                               bg=bg, bold=is_exceed, size_heb=9, size_eng=7)
                    if n_rows_for_drill > 1:
                        tc   = table.cell(row_idx, ci)._tc
                        tcPr = tc.get_or_add_tcPr()
                        for old_vm in tcPr.findall(qn('w:vMerge')): tcPr.remove(old_vm)
                        vm = OxmlElement('w:vMerge')
                        if seen == 0:
                            vm.set(qn('w:val'), 'restart')
                        tcPr.append(vm)
                else:
                    write_cell(table.cell(row_idx, ci), str(val) if val is not None else "",
                               bg=bg, bold=is_exceed, size_heb=9, size_eng=7)

        # ── מקרא ─────────────────────────────────────────────────────────────
        lp = doc.add_paragraph()
        lp.paragraph_format.space_before = Pt(4)
        lp.paragraph_format.space_after  = Pt(0)
        lp_pPr = lp._p.get_or_add_pPr()
        for old_jc in lp_pPr.findall(_w('jc')): lp_pPr.remove(old_jc)
        for old_bi in lp_pPr.findall(_w('bidi')): lp_pPr.remove(old_bi)
        jc_lp = _lxml.SubElement(lp_pPr, _w('jc')); jc_lp.set(_w('val'), 'right')

        _W2 = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
        def _w2(tag): return f'{{{_W2}}}{tag}'

        def leg_word_run(para, word, hex_color):
            r_el   = _lxml.SubElement(para._p, _w2('r'))
            rPr_el = _lxml.SubElement(r_el,    _w2('rPr'))
            rF = _lxml.SubElement(rPr_el, _w2('rFonts'))
            for a in ('ascii','hAnsi','cs','eastAsia'): rF.set(_w2(a), 'David')
            b_el = _lxml.SubElement(rPr_el, _w2('b')); b_el.set(_w2('val'), '1')
            bC   = _lxml.SubElement(rPr_el, _w2('bCs')); bC.set(_w2('val'), '1')
            sz   = _lxml.SubElement(rPr_el, _w2('sz'));   sz.set(_w2('val'), '18')
            szC  = _lxml.SubElement(rPr_el, _w2('szCs')); szC.set(_w2('val'), '18')
            shd  = _lxml.SubElement(rPr_el, _w2('shd'))
            shd.set(_w2('val'), 'clear'); shd.set(_w2('color'), 'auto')
            shd.set(_w2('fill'), hex_color)
            _lxml.SubElement(rPr_el, _w2('rtl'))
            t = _lxml.SubElement(r_el, _w2('t'))
            t.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
            t.text = word

        def leg_plain_run(para, text):
            r_el   = _lxml.SubElement(para._p, _w2('r'))
            rPr_el = _lxml.SubElement(r_el,    _w2('rPr'))
            rF = _lxml.SubElement(rPr_el, _w2('rFonts'))
            for a in ('ascii','hAnsi','cs','eastAsia'): rF.set(_w2(a), 'David')
            sz  = _lxml.SubElement(rPr_el, _w2('sz'));   sz.set(_w2('val'), '18')
            szC = _lxml.SubElement(rPr_el, _w2('szCs')); szC.set(_w2('val'), '18')
            _lxml.SubElement(rPr_el, _w2('rtl'))
            t = _lxml.SubElement(r_el, _w2('t'))
            t.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
            t.text = text

        if has_yel:
            leg_word_run(lp, "בצהוב", "FFFF00")
            leg_plain_run(lp, " - חריגה מערך הסף VSL")
        if has_yel and has_org:
            leg_plain_run(lp, "     ")
        if has_org:
            leg_word_run(lp, "בכתום", "FFC000")
            leg_plain_run(lp, " - חריגה מערך הסף TIER 1")

    buf = io.BytesIO(); doc.save(buf); buf.seek(0)
    return buf.getvalue()



def build_generic_transposed_word(xl_file_bytes, table_num, title_suffix,
                                   n_info_cols, n_hdr_rows,
                                   hdr_color="8DB4E2", group_color="DCE6F1",
                                   page_size="A3", landscape=True):
    """
    בונה Word לטבלאות transposed (שורות=תרכובות, עמודות=קידוחים).
    קורא ישירות מהאקסל ושומר merges, צבעים ומבנה.
    """
    import io, re
    from lxml import etree as _lxml
    from docx import Document
    from docx.shared import Pt, Emu, Twips
    from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
    from docx.enum.section import WD_ORIENT
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    import openpyxl

    YELLOW = "FFFF00"
    ORANGE = "FFC000"
    WHITE  = "FFFFFF"
    HDR_BG = hdr_color
    GRP_BG = group_color

    _W = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
    def _w(tag): return f'{{{_W}}}{tag}'

    wb = openpyxl.load_workbook(io.BytesIO(xl_file_bytes), data_only=True)
    ws = wb.active
    n_rows = ws.max_row
    n_cols = ws.max_column

    # ── מיפוי מיזוגים ─────────────────────────────────────────────────────────
    # master_map[(r,c)] = (min_r,min_c,max_r,max_c)
    # skip_set = תאים שאינם master
    master_map = {}
    skip_set   = set()
    for rng in ws.merged_cells.ranges:
        for r in range(rng.min_row, rng.max_row+1):
            for c in range(rng.min_col, rng.max_col+1):
                master_map[(r,c)] = (rng.min_row, rng.min_col, rng.max_row, rng.max_col)
                if r != rng.min_row or c != rng.min_col:
                    skip_set.add((r,c))

    def xl_bg(cell):
        """מחזיר צבע רקע תא Excel כ-hex 6 תווים, None אם לבן/שקוף"""
        try:
            fill = cell.fill
            if fill and fill.fill_type not in (None, 'none'):
                fg = fill.fgColor
                if fg.type == 'rgb':
                    h = fg.rgb[-6:].upper()
                    if h not in ('FFFFFF','000000',''):
                        return h
        except: pass
        return None

    def is_hdr_color(h):
        return h and h.upper() in (HDR_BG.upper(), '8DB4E2','0070C0','2E75B6',
                                    'B7D7F0','BDD7EE','9DC3E6','DAEEF3',
                                    '4472C4','00B0F0', '70AD47')

    def is_grp_color(h):
        return h and h.upper() in (GRP_BG.upper(), 'DCE6F1','D9E2F3','EBF0FA')

    def has_heb(text):
        t = str(text).strip()
        if not t: return False
        if any('\u05d0' <= c <= '\u05ea' for c in t): return True
        if any(c.isalpha() for c in t): return False
        return True

    # ── גודל דף ────────────────────────────────────────────────────────────────
    PAGE = {
        "A3":      (int(16.54*914400), int(11.69*914400)),
        "A4":      (int(11.69*914400), int(8.27*914400)),
        "Tabloid": (int(17*914400),    int(11*914400)),
    }
    pw, ph = PAGE.get(page_size, PAGE["A3"])
    page_w = max(pw,ph) if landscape else min(pw,ph)
    page_h = min(pw,ph) if landscape else max(pw,ph)
    MARGIN    = int(0.4 * 914400)
    content_w = page_w - 2 * MARGIN

    def set_page(section):
        section.orientation  = WD_ORIENT.LANDSCAPE if landscape else WD_ORIENT.PORTRAIT
        section.page_width   = page_w
        section.page_height  = page_h
        for a in ('left_margin','right_margin','top_margin','bottom_margin'):
            setattr(section, a, MARGIN)

    # ── עזרי Word ──────────────────────────────────────────────────────────────
    def set_bg(wc, hex_color):
        tcPr = wc._tc.get_or_add_tcPr()
        for old in tcPr.findall(qn('w:shd')): tcPr.remove(old)
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'),'clear'); shd.set(qn('w:color'),'auto')
        shd.set(qn('w:fill'), hex_color)
        tcPr.append(shd)

    def set_brd(wc):
        tcPr = wc._tc.get_or_add_tcPr()
        for old in tcPr.findall(qn('w:tcBorders')): tcPr.remove(old)
        brd = OxmlElement('w:tcBorders')
        for side in ('top','left','bottom','right'):
            b = OxmlElement(f'w:{side}')
            b.set(qn('w:val'),'single'); b.set(qn('w:sz'),'4')
            b.set(qn('w:space'),'0'); b.set(qn('w:color'),'000000')
            brd.append(b)
        tcPr.append(brd)

    def set_row_h(tr, h):
        trPr = tr.find(qn('w:trPr'))
        if trPr is None:
            trPr = OxmlElement('w:trPr'); tr.insert(0,trPr)
        for old in trPr.findall(qn('w:trHeight')): trPr.remove(old)
        trH = OxmlElement('w:trHeight')
        trH.set(qn('w:val'), str(h)); trH.set(qn('w:hRule'),'exact')
        trPr.append(trH)

    def set_vm(wc, restart=False):
        tc = wc._tc; tcPr = tc.get_or_add_tcPr()
        for old in tcPr.findall(qn('w:vMerge')): tcPr.remove(old)
        vm = OxmlElement('w:vMerge')
        if restart: vm.set(qn('w:val'),'restart')
        tcPr.append(vm)

    def write_cell(wc, text, bold=False, bg=WHITE, sz=16):
        """sz in half-points: 16=8pt, 18=9pt"""
        wc.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        set_bg(wc, bg); set_brd(wc)
        p = wc.paragraphs[0]; p.clear()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after  = Pt(0)
        pPr = p._p.get_or_add_pPr()
        pPr.append(OxmlElement('w:bidi'))
        jc = OxmlElement('w:jc'); jc.set(qn('w:val'),'center'); pPr.append(jc)
        if not str(text).strip(): return
        lines = str(text).split('\n')
        use_heb = has_heb(text)
        fname   = 'David' if use_heb else 'Times New Roman'
        # כתב לבן על כהה
        text_color = 'FFFFFF' if hdr_color.upper() in ('0070C0','2E75B6','4472C4') and bg == HDR_BG else None

        def _run(txt):
            r = _lxml.SubElement(p._p, _w('r'))
            rPr = _lxml.SubElement(r, _w('rPr'))
            rF = _lxml.SubElement(rPr, _w('rFonts'))
            for a in ('ascii','hAnsi','cs','eastAsia'): rF.set(_w(a), fname)
            if bold:
                _lxml.SubElement(rPr, _w('b')).set(_w('val'),'1')
                _lxml.SubElement(rPr, _w('bCs')).set(_w('val'),'1')
            if text_color:
                _lxml.SubElement(rPr, _w('color')).set(_w('val'), text_color)
            _lxml.SubElement(rPr, _w('sz')).set(_w('val'), str(sz))
            _lxml.SubElement(rPr, _w('szCs')).set(_w('val'), str(sz))
            _lxml.SubElement(rPr, _w('rtl'))
            t = _lxml.SubElement(r, _w('t'))
            t.set('{http://www.w3.org/XML/1998/namespace}space','preserve')
            t.text = txt

        def _br():
            _lxml.SubElement(_lxml.SubElement(p._p, _w('r')), _w('br'))

        for i, line in enumerate(lines):
            if i > 0: _br()
            _run(line)

    # ── רוחב עמודות ────────────────────────────────────────────────────────────
    # info cols (0..n_info_cols-1): רחבות יותר
    # drill cols: שאר
    n_drill_cols = n_cols - n_info_cols
    # חלוקה: info = 35% מסך, drill = 65%
    info_total  = int(content_w * 0.38)
    drill_total = content_w - info_total
    info_w  = info_total // n_info_cols if n_info_cols > 0 else 0
    drill_w = max(int(drill_total // max(1, n_drill_cols)), 1)

    col_ws_all = [info_w]*n_info_cols + [drill_w]*n_drill_cols

    # ── חלוקה לדפים ────────────────────────────────────────────────────────────
    drills_pp = max(5, (content_w - info_total) // max(1, drill_w))
    drill_col_idxs = list(range(n_info_cols, n_cols))
    pages = []
    i = 0
    while i < len(drill_col_idxs):
        pages.append(drill_col_idxs[i:i+drills_pp])
        i += drills_pp
    if not pages: pages = [[]]
    total_parts = len(pages)

    # ── בנה Document ───────────────────────────────────────────────────────────
    doc = Document()
    _WD = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
    styles = doc.element.find(f'{{{_WD}}}styles')
    if styles is not None:
        for b in styles.findall(f'.//{{{_WD}}}b[@{{{_WD}}}val="0"]'):
            b.getparent().remove(b)
        for b in styles.findall(f'.//{{{_WD}}}bCs[@{{{_WD}}}val="0"]'):
            b.getparent().remove(b)
    for p in doc.paragraphs:
        p._element.getparent().remove(p._element)

    def _title_run(parent_p, text, is_eng):
        fname = 'Times New Roman' if is_eng else 'David'
        r = _lxml.SubElement(parent_p, _w('r'))
        rPr = _lxml.SubElement(r, _w('rPr'))
        rF = _lxml.SubElement(rPr, _w('rFonts'))
        for a in ('ascii','hAnsi','cs','eastAsia'): rF.set(_w(a), fname)
        _lxml.SubElement(rPr, _w('b')).set(_w('val'),'1')
        _lxml.SubElement(rPr, _w('bCs')).set(_w('val'),'1')
        _lxml.SubElement(rPr, _w('u')).set(_w('val'),'single')
        _lxml.SubElement(rPr, _w('sz')).set(_w('val'), '22' if is_eng else '26')
        _lxml.SubElement(rPr, _w('szCs')).set(_w('val'), '22' if is_eng else '26')
        _lxml.SubElement(rPr, _w('rtl'))
        t = _lxml.SubElement(r, _w('t'))
        t.set('{http://www.w3.org/XML/1998/namespace}space','preserve')
        t.text = text

    for part_idx, page_drills in enumerate(pages):
        pn  = part_idx + 1
        pts = f"(חלק {pn} מתוך {total_parts})" if total_parts > 1 else ""
        title = f"טבלה מספר {table_num} {pts} – {title_suffix}".strip()

        section = doc.sections[0] if part_idx == 0 else doc.add_section()
        set_page(section)

        # כותרת
        tp = doc.add_paragraph()
        tp.paragraph_format.space_before = Pt(0)
        tp.paragraph_format.space_after  = Pt(0)
        tp.paragraph_format.keep_with_next = True
        pPr = tp._p.get_or_add_pPr()
        _lxml.SubElement(pPr, _w('bidi'))
        _lxml.SubElement(pPr, _w('jc')).set(_w('val'),'center')
        for seg, is_eng in [(p, bool(re.match(r'^[A-Za-z]+$',p)))
                            for p in re.split(r'([A-Za-z]+)', title) if p]:
            _title_run(tp._p, seg, is_eng)

        # רווח 1.5 שורות
        sp = doc.add_paragraph()
        sp.paragraph_format.space_before = Pt(0)
        sp.paragraph_format.space_after  = Pt(0)
        sp.paragraph_format.line_spacing = Pt(19.5)
        sp.paragraph_format.keep_with_next = True

        # עמודות בדף זה
        page_col_idxs = list(range(n_info_cols)) + page_drills  # 0-indexed
        n_page_cols = len(page_col_idxs)

        # טבלה
        table = doc.add_table(rows=n_rows, cols=n_page_cols)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.style = 'Table Grid'
        tblPr = table._tbl.find(qn('w:tblPr'))
        if tblPr is None:
            tblPr = OxmlElement('w:tblPr'); table._tbl.insert(0,tblPr)
        _lxml.SubElement(tblPr, _w('bidiVisual'))

        # רוחב עמודות
        for new_ci, orig_ci in enumerate(page_col_idxs):
            w = col_ws_all[orig_ci]
            for wc in table.columns[new_ci].cells:
                wc.width = Emu(w)

        # גובה שורות
        for ri in range(n_rows):
            set_row_h(table.rows[ri]._tr, 360 if ri < n_hdr_rows else 280)

        # ── מלא טבלה ──────────────────────────────────────────────────────────
        has_yel = False; has_org = False

        # בנה מפת תאים שכבר טופלו (hMerge) כדי לא לכתוב פעמיים
        hmerge_done = set()

        for new_ri in range(n_rows):
            orig_ri = new_ri + 1  # 1-indexed openpyxl

            for new_ci, orig_ci_0 in enumerate(page_col_idxs):
                orig_ci = orig_ci_0 + 1  # 1-indexed

                if (new_ri, new_ci) in hmerge_done:
                    continue

                # skip non-master merge cells (מקורי)
                if (orig_ri, orig_ci) in skip_set:
                    continue

                xl_c = ws.cell(orig_ri, orig_ci)
                val  = xl_c.value
                val_s = str(val).strip() if val is not None else ''

                # צבע רקע
                xl_fill = xl_bg(xl_c)
                is_hdr  = orig_ri <= n_hdr_rows
                is_grp  = not is_hdr and is_grp_color(xl_fill)

                if is_hdr:
                    bg = HDR_BG
                elif xl_fill == YELLOW:
                    bg = YELLOW; has_yel = True
                elif xl_fill == ORANGE:
                    bg = ORANGE; has_org = True
                elif is_grp:
                    bg = GRP_BG
                else:
                    bg = WHITE

                is_exceed = bg in (YELLOW, ORANGE)
                bold = is_hdr or is_grp or is_exceed

                wc = table.cell(new_ri, new_ci)
                write_cell(wc, val_s, bold=bold, bg=bg, sz=18 if is_hdr else 16)

                # ── vMerge: info cols בשורות header ────────────────────────────
                # cols 0..n_info_cols-1, rows 0..n_hdr_rows-1: vMerge restart בשורה 0
                if orig_ci_0 < n_info_cols and orig_ri <= n_hdr_rows:
                    set_vm(wc, restart=(orig_ri == 1))

                # ── hMerge (gridSpan) לשמות קידוחים ────────────────────────────
                mi = master_map.get((orig_ri, orig_ci))
                if mi:
                    mr1, mc1, mr2, mc2 = mi
                    # מצא כמה עמודות בדף מכוסות ע"י המיזוג
                    covered = []
                    for span_c in range(mc1, mc2+1):
                        span_c_0 = span_c - 1  # 0-indexed
                        if span_c_0 in page_col_idxs:
                            covered.append(page_col_idxs.index(span_c_0))
                    if len(covered) > 1:
                        end_new_ci = max(covered)
                        try:
                            end_wc = table.cell(new_ri, end_new_ci)
                            wc.merge(end_wc)
                            for mc in range(new_ci, end_new_ci+1):
                                hmerge_done.add((new_ri, mc))
                        except: pass

        # ── מקרא ──────────────────────────────────────────────────────────────
        lp = doc.add_paragraph()
        lp.paragraph_format.space_before = Pt(4)
        lp.paragraph_format.space_after  = Pt(0)
        lp_pPr = lp._p.get_or_add_pPr()
        for old in lp_pPr.findall(_w('jc')): lp_pPr.remove(old)
        for old in lp_pPr.findall(_w('bidi')): lp_pPr.remove(old)
        _lxml.SubElement(lp_pPr, _w('jc')).set(_w('val'),'right')

        def _leg_colored(word, color):
            r = _lxml.SubElement(lp._p, _w('r'))
            rPr = _lxml.SubElement(r, _w('rPr'))
            rF = _lxml.SubElement(rPr, _w('rFonts'))
            for a in ('ascii','hAnsi','cs','eastAsia'): rF.set(_w(a),'David')
            _lxml.SubElement(rPr, _w('b')).set(_w('val'),'1')
            _lxml.SubElement(rPr, _w('bCs')).set(_w('val'),'1')
            _lxml.SubElement(rPr, _w('sz')).set(_w('val'),'18')
            _lxml.SubElement(rPr, _w('szCs')).set(_w('val'),'18')
            shd = _lxml.SubElement(rPr, _w('shd'))
            shd.set(_w('val'),'clear'); shd.set(_w('color'),'auto'); shd.set(_w('fill'),color)
            _lxml.SubElement(rPr, _w('rtl'))
            t = _lxml.SubElement(r, _w('t'))
            t.set('{http://www.w3.org/XML/1998/namespace}space','preserve')
            t.text = word

        def _leg_plain(text):
            r = _lxml.SubElement(lp._p, _w('r'))
            rPr = _lxml.SubElement(r, _w('rPr'))
            rF = _lxml.SubElement(rPr, _w('rFonts'))
            for a in ('ascii','hAnsi','cs','eastAsia'): rF.set(_w(a),'David')
            _lxml.SubElement(rPr, _w('sz')).set(_w('val'),'18')
            _lxml.SubElement(rPr, _w('szCs')).set(_w('val'),'18')
            _lxml.SubElement(rPr, _w('rtl'))
            t = _lxml.SubElement(r, _w('t'))
            t.set('{http://www.w3.org/XML/1998/namespace}space','preserve')
            t.text = text

        if has_yel:
            _leg_colored("בצהוב", "FFFF00"); _leg_plain(" - חריגה מערך הסף VSL")
        if has_yel and has_org:
            _leg_plain("     ")
        if has_org:
            _leg_colored("בכתום", "FFC000"); _leg_plain(" - חריגה מערך הסף TIER 1")

    buf = io.BytesIO(); doc.save(buf); buf.seek(0)
    return buf.getvalue()


def build_voc_word(xl_file_bytes, table_num, page_size="A3", landscape=True):
    return build_generic_transposed_word(
        xl_file_bytes, table_num,
        title_suffix="תוצאות אנליזות VOC ו-SVOC",
        n_info_cols=7, n_hdr_rows=2,
        hdr_color="8DB4E2", group_color="DCE6F1",
        page_size=page_size, landscape=landscape
    )


def build_pfas_word(xl_file_bytes, table_num, page_size="A3", landscape=True):
    """PFAS Word - זהה לסגנון מתכות/TPH"""
    import io, re
    from lxml import etree as _lxml
    from docx import Document
    from docx.shared import Pt, Inches, Emu
    from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
    from docx.enum.section import WD_ORIENT
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    import openpyxl

    HDR_BG = "B7D7F0"   # אותו כחול כמו מתכות/TPH
    WHITE  = "FFFFFF"
    YELLOW = "FFFF00"
    ORANGE = "FFC000"
    _W = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
    def _w(tag): return f'{{{_W}}}{tag}'

    SZ_HEB = 9   # David 13pt → half-points = 26, אבל sz ב-Word = half-pt → 13pt=26
    SZ_ENG = 7   # Times New Roman 11pt → 22 half-pt... תיקון:
    # David 13pt = sz 26, Times New Roman 11pt = sz 22
    SZ_HEB_HP = 26   # David 13pt
    SZ_ENG_HP = 22   # Times New Roman 11pt

    def clean_drill_name(name):
        """S8-1.0 → S-8 | S39 (3.0) → S-39 | S43 (1.0) DUP → S-43 DUP"""
        s = str(name).strip()
        s = re.sub(r'\s*\(\d+\.?\d*\)', '', s)          # הסר (3.0)
        s = re.sub(r'-\d+\.?\d*(\s+|$)', r'\1', s)      # הסר -1.0 בסוף
        s = re.sub(r'^([A-Za-z]+)(\d+)', r'\1-\2', s)   # הוסף מקף S8→S-8
        return s.strip()

    # ── קרא אקסל ──────────────────────────────────────────────────────────────
    wb = openpyxl.load_workbook(io.BytesIO(xl_file_bytes), data_only=True)
    ws = wb.active
    n_xl_rows = ws.max_row
    n_xl_cols = ws.max_column

    def cv(r, c):
        v = ws.cell(r, c).value
        return str(v).strip() if v is not None else ''

    def cell_bg(r, c):
        try:
            fill = ws.cell(r, c).fill
            if fill and fill.fill_type not in (None, 'none'):
                fg = fill.fgColor
                if fg.type == 'rgb':
                    h = fg.rgb[-6:].upper()
                    if h == YELLOW: return YELLOW
        except: pass
        return None

    def fmt_val(raw):
        if raw is None: return ''
        s = str(raw).strip()
        if not s or s == '-': return s
        if s.startswith('<'): return s
        try:
            f = float(s)
            if f == int(f) and abs(f) >= 1000: return f'{int(f):,}'
            elif abs(f) >= 1000: return f'{f:,.1f}'
            else: return s
        except: return s

    def has_heb(text):
        t = str(text).strip()
        if not t: return False
        if any('\u05d0' <= c <= '\u05ea' for c in t): return True
        if any(c.isalpha() for c in t): return False
        return True  # מספרים → David

    DRILL_START = 7

    drills = []
    for c in range(DRILL_START, n_xl_cols + 1):
        raw_name = cv(1, c)
        depth    = cv(2, c)
        if raw_name:
            drills.append((clean_drill_name(raw_name), depth))

    compounds = []
    for r in range(3, n_xl_rows + 1):
        name = cv(r, 1)
        if not name: continue
        values = []
        for c in range(DRILL_START, n_xl_cols + 1):
            if c - DRILL_START < len(drills):
                values.append((fmt_val(ws.cell(r, c).value), cell_bg(r, c)))
        compounds.append({
            'name': name, 'cas': cv(r, 2),
            'threshold': fmt_val(cv(r, 3)), 'units': cv(r, 4),
            'lor': fmt_val(cv(r, 5)), 'values': values
        })

    if not compounds or not drills:
        raise ValueError("לא נמצאו נתונים בקובץ PFAS")

    n_compounds = len(compounds)
    n_drills    = len(drills)

    # ── גודל דף ───────────────────────────────────────────────────────────────
    def set_page(section):
        if landscape:
            section.orientation = WD_ORIENT.LANDSCAPE
            section.page_width  = Inches(16.54)
            section.page_height = Inches(11.69)
        else:
            section.orientation = WD_ORIENT.PORTRAIT
            section.page_width  = Inches(11.69)
            section.page_height = Inches(16.54)
        for a in ('left_margin','right_margin','top_margin','bottom_margin'):
            setattr(section, a, Inches(0.4))

    # ── רוחב עמודות ───────────────────────────────────────────────────────────
    usable_w = int((16.54 - 0.8) * 914400)

    # עמודת שם תרכובת רחבה מספיק ל-52 תווים ב-11pt (~4.5 אינץ')
    INFO_W = [
        int(4.50 * 914400),  # שם תרכובת
        int(1.10 * 914400),  # CAS - רחב לכל המספרים
        int(0.70 * 914400),  # ערך סף
        int(0.55 * 914400),  # יחידות
        int(0.55 * 914400),  # LOR
        int(0.55 * 914400),  # label: שם הקידוח / עומק
    ]
    N_INFO      = len(INFO_W)
    info_total  = sum(INFO_W)
    avail_drill = usable_w - info_total
    # מקסימום 8 קידוחים לדף - כמו הקובץ המקורי
    drills_per_page = 8
    drill_w         = int(avail_drill // drills_per_page)

    pages       = [drills[i:i+drills_per_page] for i in range(0, n_drills, drills_per_page)]
    total_pages = len(pages)

    # ── עזרי Word ──────────────────────────────────────────────────────────────
    def set_bg(cell, hex_color):
        tcPr = cell._tc.get_or_add_tcPr()
        for old in tcPr.findall(qn('w:shd')): tcPr.remove(old)
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear'); shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'), hex_color)
        tcPr.append(shd)

    def set_brd(cell):
        tcPr = cell._tc.get_or_add_tcPr()
        for old in tcPr.findall(qn('w:tcBorders')): tcPr.remove(old)
        brd = OxmlElement('w:tcBorders')
        for side in ('top','left','bottom','right'):
            b = OxmlElement(f'w:{side}')
            b.set(qn('w:val'), 'single'); b.set(qn('w:sz'), '4')
            b.set(qn('w:space'), '0'); b.set(qn('w:color'), '000000')
            brd.append(b)
        tcPr.append(brd)

    def set_row_h(tr, twips, rule='atLeast'):
        """atLeast = שורה לא תהיה קטנה מ-X אבל יכולה לגדול אם תוכן גדול"""
        trPr = tr.find(qn('w:trPr'))
        if trPr is None:
            trPr = OxmlElement('w:trPr'); tr.insert(0, trPr)
        for old in trPr.findall(qn('w:trHeight')): trPr.remove(old)
        trH = OxmlElement('w:trHeight')
        trH.set(qn('w:val'), str(twips))
        trH.set(qn('w:hRule'), rule)
        trPr.append(trH)

    def set_vmerge(cell, restart=False):
        tcPr = cell._tc.get_or_add_tcPr()
        for old in tcPr.findall(qn('w:vMerge')): tcPr.remove(old)
        vm = OxmlElement('w:vMerge')
        if restart: vm.set(qn('w:val'), 'restart')
        tcPr.append(vm)

    def add_run(para, text, bold=False, hp=None, color=None, underline=False):
        """hp = half-points מפורש. אם None → מחשב לפי has_heb"""
        if not str(text or '').strip(): return
        lines = str(text).split('\n')
        use_heb = has_heb(text)
        fname   = 'David' if use_heb else 'Times New Roman'
        fsize   = str(hp) if hp else (str(SZ_HEB_HP) if use_heb else str(SZ_ENG_HP))

        def _run(txt):
            r_el = _lxml.SubElement(para._p, _w('r'))
            rPr  = _lxml.SubElement(r_el,  _w('rPr'))
            rF   = _lxml.SubElement(rPr,   _w('rFonts'))
            for a in ('ascii','hAnsi','cs','eastAsia'): rF.set(_w(a), fname)
            if bold:
                _lxml.SubElement(rPr, _w('b')).set(_w('val'),   '1')
                _lxml.SubElement(rPr, _w('bCs')).set(_w('val'), '1')
            if underline:
                _lxml.SubElement(rPr, _w('u')).set(_w('val'), 'single')
            if color:
                _lxml.SubElement(rPr, _w('color')).set(_w('val'), color)
            _lxml.SubElement(rPr, _w('sz')).set(_w('val'),   fsize)
            _lxml.SubElement(rPr, _w('szCs')).set(_w('val'), fsize)
            _lxml.SubElement(rPr, _w('rtl'))
            t_el = _lxml.SubElement(r_el, _w('t'))
            t_el.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
            t_el.text = txt

        def _br():
            _lxml.SubElement(_lxml.SubElement(para._p, _w('r')), _w('br'))

        for i, line in enumerate(lines):
            if i > 0: _br()
            _run(line)

    def write_cell(cell, text, bold=False, bg=WHITE):
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        set_bg(cell, bg); set_brd(cell)
        p = cell.paragraphs[0]; p.clear()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after  = Pt(0)
        pPr = p._p.get_or_add_pPr()
        pPr.append(OxmlElement('w:bidi'))
        jc = OxmlElement('w:jc'); jc.set(qn('w:val'), 'center'); pPr.append(jc)
        if not str(text or '').strip(): return
        clr = '000000' if bg != HDR_BG else None  # שחור לתאי data
        add_run(p, str(text), bold=bold, color=clr)

    def title_segments(txt):
        parts = re.split(r'([A-Za-z]+)', txt)
        return [(p, bool(re.match(r'^[A-Za-z]+$', p))) for p in parts if p]

    def add_title_run(para_p, text, is_eng):
        fname = 'Times New Roman' if is_eng else 'David'
        fsize = '22' if is_eng else '26'
        r_el = _lxml.SubElement(para_p, _w('r'))
        rPr  = _lxml.SubElement(r_el,  _w('rPr'))
        rF   = _lxml.SubElement(rPr,   _w('rFonts'))
        for a in ('ascii','hAnsi','cs','eastAsia'): rF.set(_w(a), fname)
        _lxml.SubElement(rPr, _w('b')).set(_w('val'),   '1')
        _lxml.SubElement(rPr, _w('bCs')).set(_w('val'), '1')
        _lxml.SubElement(rPr, _w('u')).set(_w('val'), 'single')
        _lxml.SubElement(rPr, _w('sz')).set(_w('val'),   fsize)
        _lxml.SubElement(rPr, _w('szCs')).set(_w('val'), fsize)
        _lxml.SubElement(rPr, _w('rtl'))
        t_el = _lxml.SubElement(r_el, _w('t'))
        t_el.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
        t_el.text = text

    # ── בנה Document ───────────────────────────────────────────────────────────
    doc = Document()
    _WD = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
    styles = doc.element.find(f'{{{_WD}}}styles')
    if styles is not None:
        for b0 in styles.findall(f'.//{{{_WD}}}b[@{{{_WD}}}val="0"]'):
            b0.getparent().remove(b0)
        for b0 in styles.findall(f'.//{{{_WD}}}bCs[@{{{_WD}}}val="0"]'):
            b0.getparent().remove(b0)
    for p in doc.paragraphs:
        p._element.getparent().remove(p._element)

    has_yel = False
    has_org  = False

    for part_idx, page_drills in enumerate(pages):
        n_page_drills = len(page_drills)
        drill_offset  = sum(len(pages[i]) for i in range(part_idx))

        section = doc.sections[0] if part_idx == 0 else doc.add_section()
        set_page(section)

        # ── כותרת ─────────────────────────────────────────────────────────────
        pts   = f"(חלק {part_idx+1} מתוך {total_pages})" if total_pages > 1 else ""
        title = f"טבלה מספר {table_num} {pts} – תוצאות אנליזות PFAS".strip()

        tp = doc.add_paragraph()
        tp.paragraph_format.space_before   = Pt(0)
        tp.paragraph_format.space_after    = Pt(0)
        tp.paragraph_format.keep_with_next = True
        pPr_t = tp._p.get_or_add_pPr()
        sp_el = _lxml.SubElement(pPr_t, _w('spacing'))
        sp_el.set(_w('before'), '0'); sp_el.set(_w('after'), '0')
        _lxml.SubElement(pPr_t, _w('keepNext'))
        _lxml.SubElement(pPr_t, _w('bidi'))
        _lxml.SubElement(pPr_t, _w('jc')).set(_w('val'), 'center')
        for seg, is_eng in title_segments(title):
            add_title_run(tp._p, seg, is_eng)

        # ── רווח 1.5 שורות ────────────────────────────────────────────────────
        sp_p = doc.add_paragraph()
        sp_p.paragraph_format.space_before   = Pt(0)
        sp_p.paragraph_format.space_after    = Pt(0)
        sp_p.paragraph_format.keep_with_next = True
        sp_p.paragraph_format.line_spacing   = Pt(19.5)

        # ── טבלה ──────────────────────────────────────────────────────────────
        n_rows = 2 + n_compounds
        n_cols = N_INFO + n_page_drills

        table = doc.add_table(rows=n_rows, cols=n_cols)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.style = 'Table Grid'
        tblPr = table._tbl.find(qn('w:tblPr'))
        if tblPr is None:
            tblPr = OxmlElement('w:tblPr'); table._tbl.insert(0, tblPr)
        _lxml.SubElement(tblPr, _w('bidiVisual'))

        for ci in range(n_cols):
            w = INFO_W[ci] if ci < N_INFO else drill_w
            for wc in table.columns[ci].cells:
                wc.width = Emu(w)

        # גובה שורות - auto כמו הקובץ המקורי (300 twips minimum, גדל לפי תוכן)
        for ri in range(n_rows):
            set_row_h(table.rows[ri]._tr, 300, rule='auto')

        # ── 2 שורות header ────────────────────────────────────────────────────
        # עמודות info (0-4): vMerge - שורה 0+1 מאוחדות
        # עמודה 5 (label): שורה 0 = "שם הקידוח", שורה 1 = "עומק" - ללא מיזוג
        HDR_INFO_LABELS = ['שם התרכובת', 'CAS', 'ערך סף', 'יחידות', 'LOR']
        for ci in range(N_INFO - 1):  # 0-4: vMerge
            write_cell(table.cell(0, ci), HDR_INFO_LABELS[ci], bold=True, bg=HDR_BG)
            write_cell(table.cell(1, ci), '',                   bold=True, bg=HDR_BG)
            set_vmerge(table.cell(0, ci), restart=True)
            set_vmerge(table.cell(1, ci), restart=False)
        # עמודה 5: שם הקידוח / עומק - 2 שורות נפרדות
        write_cell(table.cell(0, N_INFO - 1), 'שם הקידוח', bold=True, bg=HDR_BG)
        write_cell(table.cell(1, N_INFO - 1), 'עומק',      bold=True, bg=HDR_BG)

        # כתוב header קידוחים + מזג עמודות עם אותו שם
        # בנה רשימת spans: [(name, depth, start_ci, count), ...]
        drill_spans = []
        i = 0
        while i < n_page_drills:
            dname, ddepth = page_drills[i]
            # ספור כמה עמודות עוקבות עם אותו שם
            j = i + 1
            while j < n_page_drills and page_drills[j][0] == dname:
                j += 1
            drill_spans.append((dname, ddepth, N_INFO + i, j - i))
            i = j

        for (dname, ddepth, start_ci, span) in drill_spans:
            if span == 1:
                # תא בודד - שם + עומק בשורות נפרדות
                write_cell(table.cell(0, start_ci), dname,  bold=True, bg=HDR_BG)
                write_cell(table.cell(1, start_ci), ddepth, bold=True, bg=HDR_BG)
            else:
                # מיזוג אופקי של עמודות עם אותו שם - שורה 0
                write_cell(table.cell(0, start_ci), dname, bold=True, bg=HDR_BG)
                # הגדר hMerge restart על הראשון
                tcPr0 = table.cell(0, start_ci)._tc.get_or_add_tcPr()
                for old in tcPr0.findall(qn('w:hMerge')): tcPr0.remove(old)
                hm = OxmlElement('w:hMerge'); hm.set(qn('w:val'), 'restart')
                tcPr0.append(hm)
                # שאר העמודות - hMerge continue
                for k in range(1, span):
                    write_cell(table.cell(0, start_ci + k), '', bold=True, bg=HDR_BG)
                    tcPrk = table.cell(0, start_ci + k)._tc.get_or_add_tcPr()
                    for old in tcPrk.findall(qn('w:hMerge')): tcPrk.remove(old)
                    hmk = OxmlElement('w:hMerge')
                    tcPrk.append(hmk)
                # שורה 1: עומק לכל עמודה בנפרד
                for k in range(span):
                    _, dep_k = page_drills[i - span + k] if False else page_drills[drill_spans.index((dname, ddepth, start_ci, span)) * 0 + (start_ci - N_INFO) + k]
                    write_cell(table.cell(1, start_ci + k), dep_k, bold=True, bg=HDR_BG)

        # ── שורות נתונים ──────────────────────────────────────────────────────
        for ri, cmp in enumerate(compounds):
            row_idx = ri + 2
            write_cell(table.cell(row_idx, 0), cmp['name'],      bg=WHITE)
            write_cell(table.cell(row_idx, 1), cmp['cas'],       bg=WHITE)
            write_cell(table.cell(row_idx, 2), cmp['threshold'], bg=WHITE)
            write_cell(table.cell(row_idx, 3), cmp['units'],     bg=WHITE)
            write_cell(table.cell(row_idx, 4), cmp['lor'],       bg=WHITE)
            write_cell(table.cell(row_idx, 5), cmp['lor'],       bg=WHITE)

            for di in range(n_page_drills):
                gdi = drill_offset + di
                ci  = N_INFO + di
                if gdi < len(cmp['values']):
                    val, fill = cmp['values'][gdi]
                    bg = YELLOW if fill == YELLOW else ORANGE if fill == ORANGE else WHITE
                    if bg == YELLOW: has_yel = True
                    if bg == ORANGE: has_org = True
                    write_cell(table.cell(row_idx, ci), val, bold=(bg==YELLOW), bg=bg)
                else:
                    write_cell(table.cell(row_idx, ci), '', bg=WHITE)

        # ── מקרא - זהה לVOC/generic: jc=right ללא bidi, rtl על כל run ──────────
        lp = doc.add_paragraph()
        lp.paragraph_format.space_before = Pt(4)
        lp.paragraph_format.space_after  = Pt(0)
        lp_pPr = lp._p.get_or_add_pPr()
        for old in lp_pPr.findall(_w('jc')): lp_pPr.remove(old)
        for old in lp_pPr.findall(_w('bidi')): lp_pPr.remove(old)
        _lxml.SubElement(lp_pPr, _w('jc')).set(_w('val'), 'right')

        def _leg_colored(word, color):
            r = _lxml.SubElement(lp._p, _w('r'))
            rPr = _lxml.SubElement(r, _w('rPr'))
            rF = _lxml.SubElement(rPr, _w('rFonts'))
            for a in ('ascii','hAnsi','cs','eastAsia'): rF.set(_w(a), 'David')
            _lxml.SubElement(rPr, _w('b')).set(_w('val'), '1')
            _lxml.SubElement(rPr, _w('bCs')).set(_w('val'), '1')
            _lxml.SubElement(rPr, _w('sz')).set(_w('val'), '26')
            _lxml.SubElement(rPr, _w('szCs')).set(_w('val'), '26')
            shd = _lxml.SubElement(rPr, _w('shd'))
            shd.set(_w('val'), 'clear'); shd.set(_w('color'), 'auto'); shd.set(_w('fill'), color)
            _lxml.SubElement(rPr, _w('rtl'))
            t = _lxml.SubElement(r, _w('t'))
            t.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
            t.text = word

        def _leg_plain(text):
            r = _lxml.SubElement(lp._p, _w('r'))
            rPr = _lxml.SubElement(r, _w('rPr'))
            rF = _lxml.SubElement(rPr, _w('rFonts'))
            for a in ('ascii','hAnsi','cs','eastAsia'): rF.set(_w(a), 'David')
            _lxml.SubElement(rPr, _w('sz')).set(_w('val'), '26')
            _lxml.SubElement(rPr, _w('szCs')).set(_w('val'), '26')
            _lxml.SubElement(rPr, _w('rtl'))
            t = _lxml.SubElement(r, _w('t'))
            t.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
            t.text = text

        if has_yel or has_org:
            if has_yel:
                _leg_colored("בצהוב", "FFFF00"); _leg_plain(" - חריגה מערך הסף VSL")
            if has_yel and has_org:
                _leg_plain("     ")
            if has_org:
                _leg_colored("בכתום", "FFC000"); _leg_plain(" - חריגה מערך הסף TIER 1")

    buf = io.BytesIO(); doc.save(buf); buf.seek(0)
    return buf.getvalue()


with tab_word:
    st.header("📄 יצוא דוח Word")
    st.caption("העלה קובץ Excel מעובד מטאב Excel, בחר מספר טבלה וצור דוח Word")
    st.markdown("---")

    with st.expander("🛢️ טבלת TPH"):
        wt1, wt2, wt3, wt4 = st.columns([4,1,1,1])
        with wt1:
            tph_file = st.file_uploader("העלה קובץ Excel של TPH", type=["xlsx","xls"], key="wtph")
        with wt2:
            tph_num  = st.number_input("מספר טבלה", min_value=1, max_value=99, value=1, step=1, key="wtph_num")
        with wt3:
            tph_page = st.selectbox("סוג דף", ["A4","Tabloid"], key="wtph_page")
        with wt4:
            tph_land = st.selectbox("כיוון", ["לרוחב","לאורך"], key="wtph_land") == "לרוחב"
        if tph_file:
            if st.button("📄 צור דוח Word – TPH", type="primary", use_container_width=True, key="btn_tph"):
                try:
                    with st.spinner("⏳ בונה דוח..."):
                        docx_bytes = build_tph_word(tph_file.read(), int(tph_num), tph_page, tph_land)
                    st.success("✅ הדוח נוצר!")
                    st.download_button("⬇️ הורד דוח Word – TPH", data=docx_bytes,
                        file_name=f"TPH_table_{tph_num}.docx",
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                        use_container_width=True, key="dl_tph")
                except Exception as e:
                    st.error(f"❌ שגיאה: {e}")
                    import traceback; st.code(traceback.format_exc())
        else:
            st.info("👆 העלה קובץ Excel של TPH")

    with st.expander("⚗️ טבלת Metals"):
        wm1, wm2, wm3, wm4 = st.columns([4,1,1,1])
        with wm1:
            metals_file = st.file_uploader("העלה קובץ Excel של Metals", type=["xlsx","xls"], key="wmetals")
        with wm2:
            metals_num  = st.number_input("מספר טבלה", min_value=1, max_value=99, value=2, step=1, key="wmetals_num")
        with wm3:
            metals_page = st.selectbox("סוג דף", ["A3","A4","Tabloid"], key="wmetals_page")
        with wm4:
            metals_land = st.selectbox("כיוון", ["לרוחב","לאורך"], key="wmetals_land") == "לרוחב"
        if metals_file:
            if st.button("📄 צור דוח Word – Metals", type="primary", use_container_width=True, key="btn_metals"):
                try:
                    with st.spinner("⏳ בונה דוח..."):
                        docx_bytes = build_metals_word(metals_file.read(), int(metals_num), page_size=metals_page, landscape=metals_land)
                    st.success("✅ הדוח נוצר!")
                    st.download_button("⬇️ הורד דוח Word – Metals", data=docx_bytes,
                        file_name=f"Metals_table_{metals_num}.docx",
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                        use_container_width=True, key="dl_metals")
                except Exception as e:
                    st.error(f"❌ שגיאה: {e}")
                    import traceback; st.code(traceback.format_exc())
        else:
            st.info("👆 העלה קובץ Excel של Metals")

    with st.expander("🧪 טבלת VOC+SVOC"):
        wv1, wv2, wv3, wv4 = st.columns([4,1,1,1])
        with wv1:
            voc_file = st.file_uploader("העלה קובץ Excel של VOC+SVOC", type=["xlsx","xls"], key="wvoc")
        with wv2:
            voc_num  = st.number_input("מספר טבלה", min_value=1, max_value=99, value=3, step=1, key="wvoc_num")
        with wv3:
            voc_page = st.selectbox("סוג דף", ["A3","A4","Tabloid"], key="wvoc_page")
        with wv4:
            voc_land = st.selectbox("כיוון", ["לרוחב","לאורך"], key="wvoc_land") == "לרוחב"
        if voc_file:
            if st.button("📄 צור דוח Word – VOC+SVOC", type="primary", use_container_width=True, key="btn_voc"):
                try:
                    with st.spinner("⏳ בונה דוח..."):
                        docx_bytes = build_voc_word(voc_file.read(), int(voc_num), page_size=voc_page, landscape=voc_land)
                    st.success("✅ הדוח נוצר!")
                    st.download_button("⬇️ הורד דוח Word – VOC+SVOC", data=docx_bytes,
                        file_name=f"VOC_table_{voc_num}.docx",
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                        use_container_width=True, key="dl_voc")
                except Exception as e:
                    st.error(f"❌ שגיאה: {e}")
                    import traceback; st.code(traceback.format_exc())
        else:
            st.info("👆 העלה קובץ Excel של VOC+SVOC")

    with st.expander("🔬 טבלת PFAS"):
        wp1, wp2, wp3, wp4 = st.columns([4,1,1,1])
        with wp1:
            pfas_file = st.file_uploader("העלה קובץ Excel של PFAS", type=["xlsx","xls"], key="wpfas")
        with wp2:
            pfas_num  = st.number_input("מספר טבלה", min_value=1, max_value=99, value=4, step=1, key="wpfas_num")
        with wp3:
            pfas_page = st.selectbox("סוג דף", ["A3","A4","Tabloid"], key="wpfas_page")
        with wp4:
            pfas_land = st.selectbox("כיוון", ["לרוחב","לאורך"], key="wpfas_land") == "לרוחב"
        if pfas_file:
            if st.button("📄 צור דוח Word – PFAS", type="primary", use_container_width=True, key="btn_pfas"):
                try:
                    with st.spinner("⏳ בונה דוח..."):
                        docx_bytes = build_pfas_word(pfas_file.read(), int(pfas_num), page_size=pfas_page, landscape=pfas_land)
                    st.success("✅ הדוח נוצר!")
                    st.download_button("⬇️ הורד דוח Word – PFAS", data=docx_bytes,
                        file_name=f"PFAS_table_{pfas_num}.docx",
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                        use_container_width=True, key="dl_pfas")
                except Exception as e:
                    st.error(f"❌ שגיאה: {e}")
                    import traceback; st.code(traceback.format_exc())
        else:
            st.info("👆 העלה קובץ Excel של PFAS")


# ── TPH WORD EXPORT ──────────────────────────────────────────────────────────
